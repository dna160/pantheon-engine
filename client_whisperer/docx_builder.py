"""
DOCX builder for Human Whisperer conversation prep reports.
Generates a formatted Word document from the full 6-section HumanWhispererOutput.
"""
from io import BytesIO
import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x27, 0x61)
GOLD   = RGBColor(0xE8, 0xB0, 0x4B)
GREY   = RGBColor(0x60, 0x60, 0x60)
LGREY  = RGBColor(0x90, 0x90, 0x90)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GREEN  = RGBColor(0x27, 0x80, 0x5A)
AMBER  = RGBColor(0xD3, 0x87, 0x10)
TEAL   = RGBColor(0x1A, 0x6B, 0x72)
SLATE  = RGBColor(0x47, 0x55, 0x69)

FONT_BODY = "Calibri"

FIT_COLORS = {
    "TRUE_FIT":    ("27805A", "✓ TRUE FIT"),
    "PARTIAL_FIT": ("D38710", "~ PARTIAL FIT"),
    "NO_FIT":      ("C0392B", "✗ NO FIT"),
    "VERIFY_FIT":  ("1A6B72", "? VERIFY FIT"),
}

DEPTH_LABELS = {1: "Surface", 2: "Behavioral", 3: "Emotional", 4: "Identity"}


# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_paragraph_bg(p_element, hex_color: str):
    pPr = p_element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_para_border(p_element, side: str = "left",
                     color: str = "1E2761", sz: int = 24):
    """Add a single-side border to a paragraph."""
    pPr = p_element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bdr = OxmlElement(f"w:{side}")
    bdr.set(qn("w:val"), "single")
    bdr.set(qn("w:sz"), str(sz))
    bdr.set(qn("w:space"), "6")
    bdr.set(qn("w:color"), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


# ── Typography primitives ──────────────────────────────────────────────────────

def _bagian_heading(doc: Document, text: str) -> None:
    """
    BAGIAN (section) heading — navy background strip, white text, gold rule below.
    The most visually dominant element; unambiguously marks a new major section.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    _set_paragraph_bg(p._p, "1E2761")
    run = p.add_run(f"  {text.upper()}  ")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = WHITE
    run.font.name = FONT_BODY

    # Gold rule underneath
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after  = Pt(6)
    _set_paragraph_bg(rule._p, "E8B04B")
    rule.add_run("  ").font.size = Pt(3)


def _tahap_heading(doc: Document, text: str) -> None:
    """
    TAHAP (stage) heading — teal text, thick left border, clear spacing.
    Clearly subordinate to BAGIAN but still dominant enough to scan quickly.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    _set_para_border(p._p, side="left", color="1A6B72", sz=20)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    run.font.name = FONT_BODY


def _sub_heading(doc: Document, text: str,
                 color: RGBColor = None) -> None:
    """Level-3 heading — amber, used for Wawasan N, Sinyal sub-sections, etc."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = color or AMBER
    run.font.name = FONT_BODY


def _body(doc: Document, text: str,
          italic: bool = False, color: RGBColor = None,
          size: float = 11) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_BODY
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _caption(doc: Document, text: str) -> None:
    """Small italic grey line — used as section descriptors."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = FONT_BODY
    run.italic = True
    run.font.color.rgb = LGREY


def _gap(doc: Document, size: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(size)


# ── Table helpers ──────────────────────────────────────────────────────────────

def _two_col_table(doc: Document, rows: list[tuple[str, str]],
                   label_width_pct: float = 0.28) -> None:
    """
    Label-value table.
    label_width_pct — fraction of page width for the label column (default 28%).
    """
    page_w = Inches(6.3)  # 8.5in – 1.2in margins × 2
    label_w = int(page_w * label_width_pct)
    value_w = int(page_w * (1 - label_width_pct))

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"

    for i, (label, value) in enumerate(rows):
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]

        lc.width = label_w
        vc.width = value_w

        _set_cell_bg(lc, "EEF0F8")

        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        lr.font.name = FONT_BODY
        lr.font.color.rgb = NAVY
        lc.paragraphs[0].paragraph_format.space_before = Pt(2)
        lc.paragraphs[0].paragraph_format.space_after  = Pt(2)

        vr = vc.paragraphs[0].add_run(value or "—")
        vr.font.size = Pt(10.5)
        vr.font.name = FONT_BODY
        vc.paragraphs[0].paragraph_format.space_before = Pt(2)
        vc.paragraphs[0].paragraph_format.space_after  = Pt(2)

    _gap(doc, 8)


# ── Public API ─────────────────────────────────────────────────────────────────

def build_whisper_docx(
    prospect_name: str,
    strategy: dict,
    simulated_life: str,
    linkedin_url: str = "",
    instagram_url: str = "",
) -> bytes:
    """
    Build a Human Whisperer Conversation Prep document.

    Args:
        prospect_name:  Full name extracted from LinkedIn.
        strategy:       dict returned by run_human_whisperer().
        simulated_life: Formatted PANTHEON blueprint string.
        linkedin_url:   Source LinkedIn URL.
        instagram_url:  Source Instagram URL.

    Returns:
        Raw bytes of the .docx file.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover ──────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after  = Pt(4)
    tr = title.add_run("HUMAN WHISPERER")
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = NAVY
    tr.font.name = FONT_BODY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    sr = sub.add_run(f"Persiapan Pertemuan — {prospect_name}")
    sr.font.size = Pt(13)
    sr.font.color.rgb = GREY
    sr.font.name = FONT_BODY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(2)
    mr = meta.add_run(f"Dihasilkan: {datetime.date.today().strftime('%d %B %Y')}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = LGREY
    mr.font.name = FONT_BODY

    if linkedin_url or instagram_url:
        src = doc.add_paragraph()
        src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        src.paragraph_format.space_after = Pt(0)
        parts = []
        if linkedin_url:  parts.append(f"LinkedIn: {linkedin_url}")
        if instagram_url: parts.append(f"Instagram: {instagram_url}")
        xr = src.add_run(" · ".join(parts))
        xr.font.size = Pt(9)
        xr.font.color.rgb = LGREY
        xr.font.name = FONT_BODY

    _gap(doc, 10)

    # ── Sanity Check Banner ────────────────────────────────────────────────────
    fit_raw = (strategy.get("section_5_product_fit") or {}).get("fit_status", "VERIFY_FIT")
    hex_c, label = FIT_COLORS.get(fit_raw, ("1A6B72", fit_raw))

    fit_p = doc.add_paragraph()
    fit_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fit_p.paragraph_format.space_before = Pt(4)
    fit_p.paragraph_format.space_after  = Pt(6)
    fit_run = fit_p.add_run(f"  SANITY CHECK: {label}  ")
    fit_run.bold = True
    fit_run.font.size = Pt(12)
    fit_run.font.color.rgb = WHITE
    fit_run.font.name = FONT_BODY
    _set_paragraph_bg(fit_p._p, hex_c)

    sanity_text = strategy.get("sanity_check_summary", "")
    if sanity_text:
        _body(doc, sanity_text, italic=True, color=GREY)

    _gap(doc, 6)

    # ── Plain Language Brief ───────────────────────────────────────────────────
    brief_text = strategy.get("plain_language_brief", "")
    if brief_text:
        _sub_heading(doc, "Ringkasan Bahasa Sederhana", color=NAVY)
        _body(doc, brief_text)
        _gap(doc, 6)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — HUMAN SNAPSHOT
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 1 — Potret Manusia")
    _caption(doc, "Hanya untuk internal. Jangan pernah dibagikan dengan subjek.")

    snap = strategy.get("section_1_human_snapshot") or {}
    readiness = snap.get("readiness_level", "—")
    readiness_dots = (
        "●" * int(readiness) + "○" * (5 - int(readiness))
        if isinstance(readiness, int) else str(readiness)
    )

    _two_col_table(doc, [
        ("Siapa mereka",                        snap.get("who_they_are", "—")),
        ("Bagaimana mereka melihat dirinya",     snap.get("how_they_see_themselves", "—")),
        ("Apa yang mereka inginkan",             snap.get("what_they_want", "—")),
        ("Apa yang sebenarnya mereka butuhkan",  snap.get("what_they_actually_need", "—")),
        ("Bagaimana mereka mengambil keputusan", snap.get("how_they_make_decisions", "—")),
        ("Apa yang membuat mereka percaya",      snap.get("what_makes_them_trust", "—")),
        ("Apa yang membuat mereka menutup diri", snap.get("what_makes_them_shut_down", "—")),
        ("Titik kebanggaan",                     snap.get("pride_point", "—")),
        ("Ketakutan yang sebenarnya",            snap.get("real_fear", "—")),
        (f"Tingkat kesiapan ({readiness}/5)",   readiness_dots),
        ("Satu hal untuk diingat",               snap.get("one_thing_to_remember", "—")),
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — CONVERSATION ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 2 — Arsitektur Percakapan")
    _caption(doc, "Sistem navigasi, bukan skrip. Berbebaslah menyimpang — tetapi selalu tahu di mana Anda berada di peta.")

    arch = strategy.get("section_2_conversation_architecture") or {}

    stage_keys = [
        ("stage_1_arrive",        "Tahap 1 — Tiba"),
        ("stage_2_common_ground", "Tahap 2 — Kesamaan"),
        ("stage_4_reflect",       "Tahap 4 — Refleksi"),
        ("stage_5_reframe",       "Tahap 5 — Mengubah Sudut Pandang"),
        ("stage_6_framework",     "Tahap 6 — Kerangka Kerja"),
        ("stage_7_cta",           "Tahap 7 — Panggilan Bertindak"),
    ]

    for key, label in stage_keys:
        stage = arch.get(key) or {}
        if not stage:
            continue
        dur = stage.get("duration_minutes", "")
        _tahap_heading(doc, f"{label}  ({dur})" if dur else label)
        purpose = stage.get("purpose", "")
        if purpose:
            _body(doc, f"TUJUAN: {purpose}", italic=True, color=SLATE, size=10.5)
        content = stage.get("content", "")
        if content:
            _body(doc, content, size=11)
        _gap(doc, 4)

    # ── Stage 3 — Probe Questions ──────────────────────────────────────────────
    _tahap_heading(doc, "Tahap 3 — Pertanyaan Pendalaman  (15–20 mnt)")
    _caption(doc, "Urutan: aspirasi → gesekan → konsekuensi → emosi → kepemilikan. "
                  "Selalu akui jawaban sebelum pindah ke pertanyaan berikutnya.")

    probes = arch.get("stage_3_probe") or []
    for i, q in enumerate(probes, 1):
        if isinstance(q, dict):
            q_text  = q.get("question", "")
            purpose = q.get("purpose", "")
            depth   = q.get("depth_level", "")
            follow  = q.get("open_follow_up", "")
            back    = q.get("back_out", "")
            genome  = q.get("genome_link", "")
        else:
            q_text = str(q)
            purpose = depth = follow = back = genome = ""

        depth_label = DEPTH_LABELS.get(depth, str(depth))

        # Question text
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after  = Pt(2)
        qp.paragraph_format.left_indent  = Cm(0.2)
        _set_para_border(qp._p, side="left", color="1A6B72", sz=12)
        qr = qp.add_run(f"Q{i}:  {q_text}")
        qr.bold = True
        qr.font.size = Pt(11.5)
        qr.font.name = FONT_BODY

        # Meta row
        if any([depth, genome, purpose, follow, back]):
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.space_before = Pt(0)
            meta_p.paragraph_format.space_after  = Pt(2)
            meta_p.paragraph_format.left_indent  = Cm(0.6)

            if depth or genome:
                tag_r = meta_p.add_run(
                    f"Tingkat {depth} — {depth_label}"
                    + (f"  ·  Genom: {genome}" if genome else "")
                    + "\n"
                )
                tag_r.font.size = Pt(9.5)
                tag_r.font.color.rgb = TEAL
                tag_r.font.name = FONT_BODY

            if purpose:
                pur_r = meta_p.add_run(f"Tujuan: {purpose}\n")
                pur_r.font.size = Pt(10)
                pur_r.font.color.rgb = GREY
                pur_r.font.name = FONT_BODY

            if follow:
                fol_r = meta_p.add_run(f"Jika terbuka → {follow}\n")
                fol_r.font.size = Pt(10)
                fol_r.font.color.rgb = GREEN
                fol_r.font.name = FONT_BODY

            if back:
                bak_r = meta_p.add_run(f"Jika tertutup → {back}")
                bak_r.font.size = Pt(10)
                bak_r.font.color.rgb = AMBER
                bak_r.font.name = FONT_BODY

    _gap(doc, 6)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — SIGNAL READING GUIDE
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 3 — Panduan Membaca Sinyal")
    _caption(doc, "Kalibrasi real-time. Gunakan selama percakapan.")

    sig = strategy.get("section_3_signal_reading") or {}

    open_sigs = sig.get("open_signals") or []
    if open_sigs:
        _sub_heading(doc, "✓  Sinyal Terbuka — Gali Lebih Dalam", color=GREEN)
        for s in open_sigs:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(s)
            run.font.size = Pt(11)
            run.font.name = FONT_BODY
        _gap(doc, 4)

    close_sigs = sig.get("close_signals") or []
    if close_sigs:
        _sub_heading(doc, "✗  Sinyal Tertutup — Kurangi Tekanan", color=RED)
        for s in close_sigs:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(3)
            run = p.add_run(s)
            run.font.size = Pt(11)
            run.font.name = FONT_BODY
        _gap(doc, 4)

    back_outs = sig.get("back_out_scripts") or {}
    if isinstance(back_outs, dict) and back_outs:
        _sub_heading(doc, "Skrip Menarik Diri", color=NAVY)
        for code, script in back_outs.items():
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(4)
            bp.paragraph_format.space_after  = Pt(2)
            bp.paragraph_format.left_indent  = Cm(0.3)
            br = bp.add_run(f"{code.upper()}:  ")
            br.bold = True
            br.font.size = Pt(10.5)
            br.font.color.rgb = NAVY
            br.font.name = FONT_BODY
            sr = bp.add_run(str(script))
            sr.font.size = Pt(10.5)
            sr.font.name = FONT_BODY
        _gap(doc, 6)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — PLAIN LANGUAGE TRANSLATIONS
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 4 — Panduan Terjemahan Bahasa Sederhana")
    _caption(doc, "Untuk setiap anggota tim yang akan melihat ini tanpa latar belakang PANTHEON.")

    insights = strategy.get("section_4_plain_language_guide") or []
    for i, ins in enumerate(insights, 1):
        if not isinstance(ins, dict):
            continue
        _sub_heading(doc, f"Wawasan {i}")
        _two_col_table(doc, [
            ("TEKNIS",        ins.get("technical", "—")),
            ("SEDERHANA",     ins.get("plain",     "—")),
            ("ANALOGI",       ins.get("analogy",   "—")),
            ("SATU KALIMAT",  ins.get("one_line",  "—")),
        ])

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — PRODUCT FIT & CTA SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 5 — Kecocokan Produk & Ringkasan CTA")

    pf = strategy.get("section_5_product_fit") or {}
    fit_status = pf.get("fit_status", "VERIFY_FIT")
    _, label2 = FIT_COLORS.get(fit_status, ("1A6B72", fit_status))

    rows5 = [
        ("STATUS",                    label2),
        ("Mengapa?",                  pf.get("fit_rationale",    "—")),
        ("Masalah yang diselesaikan", pf.get("pain_it_addresses","—")),
        ("Cara memperkenalkannya",    pf.get("how_to_introduce_it","—")),
        ("Batasan jujur",             pf.get("honest_limitation", "—")),
        ("Apa yang terjadi selanjutnya", pf.get("what_happens_next","—")),
    ]
    if pf.get("what_else_they_need"):
        rows5.append(("Apa lagi yang mereka butuhkan", pf["what_else_they_need"]))
    if pf.get("honest_redirect"):
        rows5.append(("Pengalihan jujur", pf["honest_redirect"]))

    _gap(doc, 4)
    _two_col_table(doc, rows5)

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — POST-CONVERSATION PROTOCOL
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Bagian 6 — Protokol Pasca-Pertemuan")
    _caption(doc, "Selesaikan dalam waktu 24 jam setelah percakapan.")

    pcp = strategy.get("section_6_post_conversation") or {}
    _two_col_table(doc, [
        ("Dalam waktu 24 jam",     pcp.get("within_24_hours",  "—")),
        ("Apa yang perlu dicatat", pcp.get("what_to_note",     "—")),
        ("Apa yang perlu diperbarui", pcp.get("what_to_update","—")),
        ("Percakapan berikutnya",  pcp.get("next_conversation","—")),
    ])

    # ═══════════════════════════════════════════════════════════════════════════
    # APPENDIX — PANTHEON LIFE BLUEPRINT (RAW)
    # ═══════════════════════════════════════════════════════════════════════════
    _bagian_heading(doc, "Appendix — PANTHEON Life Blueprint (Raw)")
    _caption(doc, "Genome scores, life layers, voice print, and mutation log.")

    bp_para = doc.add_paragraph()
    bp_para.paragraph_format.space_before = Pt(4)
    bp_run = bp_para.add_run(simulated_life)
    bp_run.font.size  = Pt(8.5)
    bp_run.font.name  = "Courier New"
    bp_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Serialize ──────────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
