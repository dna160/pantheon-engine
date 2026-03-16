import os
import re
import json
from datetime import datetime
from pathlib import Path

DIVIDER = "═" * 70

def _build_chronesthesia_directive(agent: dict) -> str:
    """Generate a cognitive mode directive based on chronesthesia_capacity score."""
    chrono = agent.get("chronesthesia_capacity", 50)
    if chrono < 25:
        return (
            "\n[COGNITIVE MODE: Present-focused. Rarely considers long-term consequences. "
            "Reacts to immediate stimuli. Decisions driven by what feels right NOW.]"
        )
    elif chrono < 50:
        return (
            "\n[COGNITIVE MODE: Moderate foresight. Can think 1-2 years ahead when prompted, "
            "but defaults to near-term. Occasionally references future plans but not systematically.]"
        )
    elif chrono < 75:
        return (
            "\n[COGNITIVE MODE: Active future simulator. Before major decisions, mentally projects "
            "outcomes 3-5 years forward. Weighs current choices against long-term identity goals. "
            "References origin/formation memories when evaluating whether choices align with core values.]"
        )
    else:
        return (
            "\n[COGNITIVE MODE: Vivid mental time traveler. Constantly simulates future scenarios. "
            "Decisions filtered through a 10-year projection. Past memories (origin/formation layers) "
            "actively queried to validate current choices. May hesitate due to overthinking consequences. "
            "Strategic, but can be paralyzed by scenario anxiety.]"
        )

def _build_agent_context(agent: dict) -> str:
    """Compact but information-dense agent profile string for prompt injection."""
    genome = (
        f"Genome (1-100): O={agent.get('openness')} C={agent.get('conscientiousness')} "
        f"E={agent.get('extraversion')} A={agent.get('agreeableness')} N={agent.get('neuroticism')} | "
        f"CommStyle={agent.get('communication_style')} Dec={agent.get('decision_making')} "
        f"Brand={agent.get('brand_relationship')} Influ={agent.get('influence_susceptibility')} "
        f"EmotExp={agent.get('emotional_expression')} Conflict={agent.get('conflict_behavior')} | "
        f"IdFusion={agent.get('identity_fusion', 50)} Chrono={agent.get('chronesthesia_capacity', 50)} "
        f"ToMSelf={agent.get('tom_self_awareness', 50)} ToMSocial={agent.get('tom_social_modeling', 50)} "
        f"ExecFlex={agent.get('executive_flexibility', 50)}"
    )

    age = agent.get("age", 30)
    if age <= 28:
        layer = agent.get("formation_layer") or agent.get("independence_layer")
    elif age <= 38:
        layer = agent.get("independence_layer") or agent.get("maturity_layer")
    elif age < 60:
        layer = agent.get("maturity_layer") or agent.get("independence_layer")
    else:
        layer = agent.get("legacy_layer") or agent.get("maturity_layer")

    layer_text = ""
    if isinstance(layer, dict):
        layer_text = (
            f"\nLife stage: {layer.get('summary', '')}"
            f"\nPsychological profile: {layer.get('psychological_impact', '')}"
        )

    chrono_text = _build_chronesthesia_directive(agent) if age < 60 else ""
    vp = agent.get("voice_print") or {}
    triggers = ""
    if isinstance(vp, dict) and vp.get("persuasion_triggers"):
        triggers = f"\nPersuasion triggers: {', '.join(vp['persuasion_triggers'])}"

    return (
        f"Age: {agent.get('age')} | Demographic: {agent.get('target_demographic')} | "
        f"Region: {agent.get('region', 'Medan, Indonesia')}\n"
        f"Culture: {agent.get('cultural_background', 'Unspecified')} | Religion: {agent.get('religion', 'Unspecified')}\n"
        f"{genome}{layer_text}{chrono_text}{triggers}"
    )

def _save_report(report: str, target: str, brief: str, client: str = "") -> tuple[Path, str]:
    """Save report as .md and .docx."""
    slug = re.sub(r"[^\w]+", "_", target).strip("_")
    client_folder = re.sub(r"[^\w]+", "_", client).strip("_") if client.strip() else "Unnamed"
    out_dir = Path(__file__).parent.parent / "reports" / client_folder
    out_dir.mkdir(parents=True, exist_ok=True)
    print(f"PANTHEON_CLIENT_FOLDER::{client_folder}")

    if client:
        client_slug = re.sub(r"[^\w]+", "_", client).strip("_")
        version = 1
        while True:
            base_name = f"PANTHEON_REPORT_{client_slug}_v{version}"
            if not ((out_dir / f"{base_name}.docx").exists() or (out_dir / f"{base_name}.md").exists()):
                break
            version += 1
    else:
        base_name = f"PANTHEON_Report_{slug}"

    docx_path = out_dir / f"{base_name}.docx"
    md_path   = out_dir / f"{base_name}.md"
    print(f"PANTHEON_OUTPUT_FILE::{base_name}")

    md_content = (
        f"# PANTHEON Research Intelligence Report\n"
        f"**Target demographic:** {target}  \n"
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  \n"
        f"**Brief:** {brief}\n\n"
        f"---\n\n"
        + report
    )
    md_path.write_text(md_content, encoding="utf-8")
    print(f"  Markdown saved  → {md_path}")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        title_para = doc.add_heading("PANTHEON Research Intelligence Report", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta = doc.add_paragraph()
        meta.add_run("Target demographic: ").bold = True
        meta.add_run(target)
        meta.add_run("\nGenerated: ").bold = True
        meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        meta.add_run("\nCampaign brief: ").bold = True
        meta.add_run(brief)
        doc.add_paragraph()

        for raw_line in report.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                doc.add_paragraph()
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
                continue
            if re.match(r"^[-=─]{3,}\s*$", line):
                doc.add_paragraph("─" * 50)
                continue
            para = doc.add_paragraph()
            segments = re.split(r"(\*\*[^*]+\*\*)", line)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    run = para.add_run(seg[2:-2])
                    run.bold = True
                else:
                    para.add_run(seg)
        doc.save(str(docx_path))
        print(f"  Word doc saved  → {docx_path}")
    except ImportError:
        print("  python-docx not installed — skipping .docx output")
    except Exception as exc:
        print(f"  Word doc failed : {exc}")

    return md_path, base_name
