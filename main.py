"""
main.py — PANTHEON Execution Engine (Modal)
Orchestrates the full 5-node serverless pipeline on Modal.

Architecture:
  Node 1  — Supabase intake & query                  (synchronous, runs locally)
  Node 2  — Runtime snapshot per agent               (parallel Haiku via .map())
  Node 3  — Phase A mass session per agent           (parallel Haiku via .starmap())
  Node 4  — Phase B breakout room debates per group  (parallel Sonnet via .starmap())
  Node 5  — Phase C synthesis final report           (single Sonnet, runs remotely)

Secrets: Modal secret named "pantheon-secrets" must contain:
  SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, ANTHROPIC_API_KEY

Usage:
  modal run main.py                              # uses defaults below
  modal run main.py --target "Medanese Upper Middle Class, 25-45" --brief "..." --limit 10
"""
import os
import json
import random
import re
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import modal

# ─────────────────────────────────────────────────────────────────────────────
# Modal App & Container Image
# ─────────────────────────────────────────────────────────────────────────────

image = (
    modal.Image.debian_slim(python_version="3.11")
    .pip_install(
        "supabase>=2.0.0",
        "anthropic>=0.50.0",
        "fastapi[standard]>=0.115.0",
        "pydantic>=2.4.0",
    )
)

app = modal.App(name="pantheon-engine")

SECRETS = [modal.Secret.from_name("pantheon-secrets")]

# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers (inlined so Modal workers receive them without local imports)
# ─────────────────────────────────────────────────────────────────────────────

DIVIDER = "═" * 70

def _build_chronesthesia_directive(agent: dict) -> str:
    """Generate a cognitive mode directive based on chronesthesia_capacity score."""
    chrono = agent.get("chronesthesia_capacity", 50)
    if chrono < 25:
        return (
            "\n[COGNITIVE MODE: Present-focused. Rarely considers long-term consequences. "
            "Reacts to immediate stimuli. Decisions driven by what feels right NOW.]"
        )
    elif chrono < 50:
        return (
            "\n[COGNITIVE MODE: Moderate foresight. Can think 1-2 years ahead when prompted, "
            "but defaults to near-term. Occasionally references future plans but not systematically.]"
        )
    elif chrono < 75:
        return (
            "\n[COGNITIVE MODE: Active future simulator. Before major decisions, mentally projects "
            "outcomes 3-5 years forward. Weighs current choices against long-term identity goals. "
            "References origin/formation memories when evaluating whether choices align with core values.]"
        )
    else:
        return (
            "\n[COGNITIVE MODE: Vivid mental time traveler. Constantly simulates future scenarios. "
            "Decisions filtered through a 10-year projection. Past memories (origin/formation layers) "
            "actively queried to validate current choices. May hesitate due to overthinking consequences. "
            "Strategic, but can be paralyzed by scenario anxiety.]"
        )


def _build_agent_context(agent: dict) -> str:
    """Compact but information-dense agent profile string for prompt injection."""
    genome = (
        f"Genome (1-100): O={agent.get('openness')} C={agent.get('conscientiousness')} "
        f"E={agent.get('extraversion')} A={agent.get('agreeableness')} N={agent.get('neuroticism')} | "
        f"CommStyle={agent.get('communication_style')} Dec={agent.get('decision_making')} "
        f"Brand={agent.get('brand_relationship')} Influ={agent.get('influence_susceptibility')} "
        f"EmotExp={agent.get('emotional_expression')} Conflict={agent.get('conflict_behavior')} | "
        f"IdFusion={agent.get('identity_fusion', 50)} Chrono={agent.get('chronesthesia_capacity', 50)} "
        f"ToMSelf={agent.get('tom_self_awareness', 50)} ToMSocial={agent.get('tom_social_modeling', 50)} "
        f"ExecFlex={agent.get('executive_flexibility', 50)}"
    )

    age = agent.get("age", 30)
    # For agents under 60, use lived life layers; legacy_layer is aspirational, not static truth
    if age <= 28:
        layer = agent.get("formation_layer") or agent.get("independence_layer")
    elif age <= 38:
        layer = agent.get("independence_layer") or agent.get("maturity_layer")
    elif age < 60:
        layer = agent.get("maturity_layer") or agent.get("independence_layer")
    else:
        layer = agent.get("legacy_layer") or agent.get("maturity_layer")

    layer_text = ""
    if isinstance(layer, dict):
        layer_text = (
            f"\nLife stage: {layer.get('summary', '')}"
            f"\nPsychological profile: {layer.get('psychological_impact', '')}"
        )

    # Chronesthesia directive replaces static legacy_layer for agents under 60
    chrono_text = _build_chronesthesia_directive(agent) if age < 60 else ""

    vp = agent.get("voice_print") or {}
    triggers = ""
    if isinstance(vp, dict) and vp.get("persuasion_triggers"):
        triggers = f"\nPersuasion triggers: {', '.join(vp['persuasion_triggers'])}"

    return (
        f"Age: {agent.get('age')} | Demographic: {agent.get('target_demographic')} | "
        f"Region: {agent.get('region', 'Medan, Indonesia')}\n"
        f"Culture: {agent.get('cultural_background', 'Unspecified')} | Religion: {agent.get('religion', 'Unspecified')}\n"
        f"{genome}{layer_text}{chrono_text}{triggers}"
    )


# ────────────────────────────────────────────────────────────────────────────
# GENESIS ENGINE  — dynamic agent seeding (ported from seed_genomes.py)
# Runs locally (same process as the local_entrypoint). No Modal container needed.
# ────────────────────────────────────────────────────────────────────────────

_GENOME_TOOL = {
    "name": "submit_genome",
    "description": "Submit all personality integer scores and voice print. ALL fields required.",
    "input_schema": {
        "type": "object",
        "properties": {
            "openness":                {"type": "integer", "description": "1=closed-minded, 100=experimental"},
            "conscientiousness":       {"type": "integer", "description": "1=impulsive, 100=methodical"},
            "extraversion":            {"type": "integer", "description": "1=introvert, 100=extrovert"},
            "agreeableness":           {"type": "integer", "description": "1=adversarial, 100=cooperative"},
            "neuroticism":             {"type": "integer", "description": "1=stable, 100=anxious"},
            "communication_style":     {"type": "integer", "description": "1=indirect, 100=direct"},
            "decision_making":         {"type": "integer", "description": "1=gut-driven, 100=data-driven"},
            "brand_relationship":      {"type": "integer", "description": "1=price-driven, 100=brand-loyal"},
            "influence_susceptibility":{"type": "integer", "description": "1=impervious, 100=susceptible"},
            "emotional_expression":    {"type": "integer", "description": "1=stoic, 100=openly emotional"},
            "conflict_behavior":       {"type": "integer", "description": "1=avoider, 100=confrontational"},
            "literacy_and_articulation": {
                "type": "integer",
                "description": "1=barely literate/inarticulate with very simple vocabulary, 100=eloquent and highly educated. Dictates vocabulary complexity, sentence structure, and speaking confidence."
            },
            "socioeconomic_friction": {
                "type": "integer",
                "description": "1=minimal barriers/comfortable life trajectory, 100=severe systemic barriers, crushing debt, career stagnation, or chronic financial precarity. Shapes pessimism, risk-aversion, and distrust of financial products."
            },
            "identity_fusion": {
                "type": "integer",
                "description": "1=pure individualist (evaluates everything on personal merit), 100=visceral group oneness (will sacrifice personal utility to protect family/clan/community honor). Drives mianzi, marga, guanxi, familismo behaviors."
            },
            "chronesthesia_capacity": {
                "type": "integer",
                "description": "1=present-only thinker (reactive, no future simulation), 100=vivid mental time traveler (constantly projects decisions 5-10 years forward, queries past memories to validate choices)."
            },
            "tom_self_awareness": {
                "type": "integer",
                "description": "1=blind to own emotional states and biases, 100=deep self-reflection (accurately identifies own motivations, recognizes own cognitive distortions)."
            },
            "tom_social_modeling": {
                "type": "integer",
                "description": "1=oblivious to how others perceive them, 100=reads rooms perfectly (accurately models what others think of them, predicts social reactions)."
            },
            "executive_flexibility": {
                "type": "integer",
                "description": "1=base traits leak into all contexts regardless of social norms, 100=can fully override impulses in context-appropriate situations (the professional who is privately anxious but appears calm)."
            },
            "religion": {
                "type": "string",
                "description": "Specific religious practice and denomination/sect, defining behavior/diet/finances (e.g., 'Conservative Sunni Muslim - halal strict', 'Devout Catholic - family-centric', 'Nominally Buddhist', 'Pentecostal - tithing')."
            },
            "cultural_background": {
                "type": "string",
                "description": "The ethno-cultural context dictating family expectations, hierarchy, and shame mechanics (e.g., 'Third-generation Chinese-Indonesian', 'Minangkabau diaspora', 'Batak Toba strictly adhering to Adat', 'Javanese')."
            },
            "voice_print": {
                "type": "object",
                "properties": {
                    "vocabulary_level":    {"type": "string"},
                    "filler_words":        {"type": "array", "items": {"type": "string"}},
                    "persuasion_triggers": {"type": "array", "items": {"type": "string"}},
                    "conflict_style":      {"type": "string"}
                },
                "required": ["vocabulary_level", "filler_words", "persuasion_triggers", "conflict_style"],
                "additionalProperties": False
            }
        },
        "required": [
            "openness", "conscientiousness", "extraversion", "agreeableness", "neuroticism",
            "communication_style", "decision_making", "brand_relationship",
            "influence_susceptibility", "emotional_expression", "conflict_behavior",
            "literacy_and_articulation", "socioeconomic_friction",
            "identity_fusion", "chronesthesia_capacity", "tom_self_awareness",
            "tom_social_modeling", "executive_flexibility",
            "religion", "cultural_background", "voice_print"
        ],
        "additionalProperties": False
    }
}

_LAYER_DEF = {
    "type": "object",
    "properties": {
        "summary":              {"type": "string", "description": "MAX 100 chars. 1 sentence."},
        "key_events":           {"type": "array", "items": {"type": "string"}, "description": "Exactly 3 items, each MAX 60 chars"},
        "psychological_impact": {"type": "string", "description": "MAX 80 chars. 1 sentence."}
    },
    "required": ["summary", "key_events", "psychological_impact"],
    "additionalProperties": False
}

_BLUEPRINT_TOOL = {
    "name": "submit_blueprint",
    "description": "Submit all 5 life stage layers. ALL layers required.",
    "input_schema": {
        "type": "object",
        "properties": {
            "origin_layer":       _LAYER_DEF,
            "formation_layer":    _LAYER_DEF,
            "independence_layer": _LAYER_DEF,
            "maturity_layer":     _LAYER_DEF,
            "legacy_layer":       _LAYER_DEF,
        },
        "required": ["origin_layer", "formation_layer", "independence_layer", "maturity_layer", "legacy_layer"],
        "additionalProperties": False
    }
}

# Trait names the mutation log is allowed to reference.
_MUTABLE_TRAITS = [
    "openness", "conscientiousness", "extraversion", "agreeableness", "neuroticism",
    "communication_style", "decision_making", "brand_relationship",
    "influence_susceptibility", "emotional_expression", "conflict_behavior",
    "literacy_and_articulation", "socioeconomic_friction",
    "identity_fusion", "chronesthesia_capacity", "tom_self_awareness",
    "tom_social_modeling", "executive_flexibility",
]

_MUTATION_TOOL = {
    "name": "submit_mutation_log",
    "description": (
        "Submit the Nature vs. Nurture mutation log. "
        "This is a variable-length array — zero events for a stable period, "
        "multiple events for a chaotic one. Do NOT force one event per stage."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "genome_mutation_log": {
                "type": "array",
                "description": (
                    "Chronological list of life events that mutated the base genome. "
                    "Can be empty [] for a completely uneventful life. "
                    "No fixed number per stage — let the life path dictate it."
                ),
                "items": {
                    "type": "object",
                    "properties": {
                        "life_stage": {
                            "type": "string",
                            "enum": ["Origin", "Formation", "Independence", "Maturity", "Legacy"],
                            "description": "Which life stage this event belongs to."
                        },
                        "event_description": {
                            "type": "string",
                            "description": "Max 80 chars. One specific event that altered this person."
                        },
                        "trait_modifiers": {
                            "type": "object",
                            "description": (
                                "Key-value pairs of trait name to integer shift (positive or negative). "
                                "Only include traits that actually changed. "
                                "Valid keys: openness, conscientiousness, extraversion, agreeableness, "
                                "neuroticism, communication_style, decision_making, brand_relationship, "
                                "influence_susceptibility, emotional_expression, conflict_behavior, "
                                "literacy_and_articulation, socioeconomic_friction, "
                                "identity_fusion, chronesthesia_capacity, tom_self_awareness, "
                                "tom_social_modeling, executive_flexibility."
                            ),
                            "additionalProperties": {"type": "integer"}
                        }
                    },
                    "required": ["life_stage", "event_description", "trait_modifiers"],
                    "additionalProperties": False
                }
            }
        },
        "required": ["genome_mutation_log"],
        "additionalProperties": False
    }
}

# City pool — drawn in Python before every LLM call to force geographic diversity.
# The chosen city is injected into the seed template AND the system prompt,
# eliminating any possibility of the model defaulting to Medan from context.
_INDONESIAN_CITIES = [
    "Jakarta",
    "Surabaya",
    "Bandung",
    "Medan",
    "Makassar",
    "Semarang",
    "Yogyakarta",
    "Palembang",
    "Bali (Denpasar)",
    "Balikpapan",
    "Malang",
    "Manado",
    "Pekanbaru",
    "Pontianak",
    "Banjarmasin",
]

# Cultural archetype seed templates.
# {age} = randomised integer, {city} = Python-chosen city (NOT left to the LLM).
# All Medan-specific landmarks removed — replaced with city-agnostic descriptors.
_SEED_TEMPLATES = [
    "Chinese-Indonesian {age}yo, third-generation trading family in {city}, "
    "conservative, status-conscious, clan association active, luxury-car aspirations.",

    "Batak Toba {age}yo professional (lawyer/doctor/engineer) now based in {city}, "
    "feminist but adat-rooted, heavy Instagram user, premium-mall shopper, skeptical of ads.",

    "Javanese {age}yo entrepreneur or startup founder based in {city}, "
    "GrabFood/TikTok/Tokopedia native, frugal personally but aggressive in business, "
    "peer-recommendation-driven.",

    "Minangkabau {age}yo female professional in {city}, evidence-driven, remittance sender, "
    "slow decision-maker with extreme conviction, anti-influencer, clinic or consultancy owner.",

    "Batak Karo / Dayak {age}yo regional manager or sales director in {city}, rose from field ops, "
    "sends kids to international school, peer-influenced, LinkedIn poster, brand-insecure.",

    "Chinese-Peranakan {age}yo in {city}, property investor or café owner, "
    "bilingual Chinese-dialect/Bahasa, mianzi-driven, early adopter of premium lifestyle brands.",

    "Sundanese/Javanese {age}yo salaried professional in {city} (banking/FMCG/government), "
    "quietly status-conscious, Shopee/Tokopedia power-user, brand-loyal to established names.",

    "Bugis {age}yo entrepreneur in {city}, maritime-trade heritage, tight community network, "
    "cash-first, distrustful of fintech debt products, drives hard bargains.",

    "Acehnese / Malay {age}yo in {city}, religiously observant, halal-compliant purchasing, "
    "rejects riba-adjacent products, strong family decision-making unit.",

    "Urban millennial {age}yo from {city}, mixed ethnicity, college-educated, remote-worker, "
    "high digital literacy, rents not owns, aspirational but cash-constrained.",
]

# Base system prompt — city-specific context is appended at call time.
_GENESIS_SYSTEM_BASE = (
    "You are PANTHEON Genesis Builder: generating psychologically coherent Indonesian consumer profiles.\n"
    "Your role is to create realistic synthetic humans for market research simulations.\n"
    "CRITICAL REALISM MANDATE: Do NOT make every agent a successful entrepreneur or high-earner. "
    "Include gig workers, unemployed individuals, burned-out middle managers, people burdened by debt, "
    "informal-sector workers, and those with failed businesses. "
    "Most people are NOT thriving — reflect this honestly.\n"
    "NATURE vs. NURTURE MUTATION SYSTEM: You must construct every agent's personality in 4 steps:\n"
    "  Step 1 (Base Nature): Establish a baseline Personality Genome (1-100) representing their "
    "genetic starting point before any life experience.\n"
    "  Step 2 (The Variable Life Path): Generate a chronological genome_mutation_log. "
    "You are NOT required to create an event for every life stage. "
    "A stage can have ZERO events if it was highly stable, or MULTIPLE events if it was chaotic. "
    "Let the agent's life unfold organically.\n"
    "  Step 3 (The Math): For every event that DOES occur in the log, apply specific +/- integer "
    "modifiers to the relevant traits (e.g., {\"neuroticism\": 15, \"conscientiousness\": -5}).\n"
    "  Step 4 (Final Genome): Apply all log modifiers to the Base Nature (cap between 1-100). "
    "This Final Genome dictates their behavior today. "
    "It is NOT the same as Base Nature unless no mutations occurred.\n"
    "Rules:\n"
    "  - Every final integer must reflect the cumulative result of nature + life events.\n"
    "  - Keep ALL text fields brief and concrete.\n"
    "  - Write in the third person.\n"
    "  - No generic pan-Indonesian descriptions — be specific to the assigned city and ethnicity.\n"
    "  - literacy_and_articulation must reflect their actual education level and environment, not "
    "    their aspiration. A dropout is 15-35. A university grad is 50-75. A postgrad is 70-90.\n"
    "  - socioeconomic_friction must reflect actual systemic barriers: debt, job loss, family "
    "    obligations, discrimination. Do not default to low friction unless genuinely warranted.\n"
    "  - identity_fusion: high for agents deeply embedded in clan/community/family structures; "
    "    low for individualists, cosmopolitans, or culturally displaced people.\n"
    "  - chronesthesia_capacity: high for strategic long-term planners; low for present-focused, "
    "    reactive personalities. Educated professionals with career plans score higher.\n"
    "  - tom_self_awareness: high for introspective, self-aware individuals; low for those who "
    "    act on impulse without understanding their own motivations.\n"
    "  - tom_social_modeling: high for socially sophisticated, face-conscious individuals; "
    "    low for those oblivious to social dynamics. High-context cultures score higher.\n"
    "  - executive_flexibility: high for disciplined professionals who can mask emotions; "
    "    low for emotionally transparent individuals whose feelings show immediately.\n"
    "CRITICAL: This is a stateless generation. You have NO memory of previous agents or workspace "
    "files. Do NOT default to Medan or any single city. The city for this agent is specified in "
    "the user prompt and you MUST ground ALL cultural references in that city only."
)


def _build_genesis_system(city: str, demographic: str) -> str:
    """Build a per-agent system prompt that forces geographic grounding."""
    return (
        f"{_GENESIS_SYSTEM_BASE}\n\n"
        f"ASSIGNED CITY FOR THIS AGENT: {city}\n"
        f"TARGET DEMOGRAPHIC: {demographic}\n"
        f"You must produce a profile that is authentically rooted in {city}. "
        f"Reference {city}-specific culture, economy, infrastructure, and social dynamics. "
        f"Do not reference any other city as the agent's home base."
    )


def _clamp_ints(data: dict) -> dict:
    """Clamp all 18 integer genome fields to the 1-100 range in-place."""
    for field in _MUTABLE_TRAITS:
        if field in data and data[field] is not None:
            data[field] = max(1, min(100, int(data[field])))
    return data


def _apply_mutations(base_genome: dict, mutation_log: list[dict]) -> dict:
    """
    Pure Python: walk the LLM-generated mutation log and apply each event's
    trait_modifiers to the base genome, clamping the result to [1, 100].

    Returns a NEW dict — base_genome is not mutated in-place.
    """
    final = dict(base_genome)  # shallow copy of scalar values
    for event in mutation_log:
        modifiers = event.get("trait_modifiers") or {}
        for trait, delta in modifiers.items():
            if trait not in _MUTABLE_TRAITS:
                continue  # ignore any hallucinated trait names
            current = final.get(trait, 50)
            final[trait] = max(1, min(100, int(current) + int(delta)))
    return final


def _genesis_call_tool(
    anthropic_client,
    tool_def: dict,
    user_message: str,
    system_prompt: str,          # per-agent system prompt (city-injected)
    max_tokens: int = 1000,
    retries: int = 3,
) -> dict | None:
    """Single tool-forced Haiku call with rate-limit retry."""
    import anthropic as _anthropic
    for attempt in range(1, retries + 1):
        try:
            response = anthropic_client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=max_tokens,
                system=system_prompt,
                tools=[tool_def],
                tool_choice={"type": "tool", "name": tool_def["name"]},
                messages=[{"role": "user", "content": user_message}]
            )
            block = next((b for b in response.content if b.type == "tool_use"), None)
            if block is None:
                raise ValueError("No tool_use block returned")
            return block.input
        except _anthropic.RateLimitError as e:
            wait = 65
            try:
                wait = int(e.response.headers.get("retry-after", 65)) + 2
            except Exception:
                pass
            print(f"    Genesis: rate limit hit — sleeping {wait}s...")
            time.sleep(wait)
        except Exception as e:
            print(f"    Genesis: attempt {attempt}/{retries} failed: {e}")
            if attempt < retries:
                time.sleep(5)
    return None


def _generate_one_agent(
    anthropic_client,
    age: int,
    seed: str,
    city: str,
    demographic: str,
) -> dict | None:
    """
    THREE Haiku calls per agent:
      1. Base genome (13 integer scores + voice_print)
      2. Mutation log (variable-length array of life events + trait modifiers)
      3. Life blueprint (5 JSONB layers)

    Python applies the mutation math via _apply_mutations() so the final
    integers stored in Supabase always reflect the environmental-altered genome.
    """
    system_prompt = _build_genesis_system(city, demographic)

    # ── CALL 1: Base Nature genome ─────────────────────────────────────────
    genome_msg = (
        f"CITY FOR THIS AGENT: {city}\n"
        f"Agent profile: {seed}\n\n"
        "STEP 1 — Base Nature: assign the raw genetic baseline for all 18 personality "
        "integers BEFORE any life experience. This is their nature before nurture. "
        f"Ground the cultural starting point in {city}."
    )
    base_genome = _genesis_call_tool(
        anthropic_client, _GENOME_TOOL, genome_msg,
        system_prompt=system_prompt, max_tokens=500,
    )
    if base_genome is None or not isinstance(base_genome.get("voice_print"), dict):
        return None
    base_genome = _clamp_ints(base_genome)

    # ── CALL 2: Mutation log (variable life events) ───────────────────────
    mutation_msg = (
        f"CITY FOR THIS AGENT: {city}\n"
        f"Agent profile: {seed}\n"
        f"Base Nature genome: "
        + ", ".join(f"{k}={base_genome.get(k)}" for k in _MUTABLE_TRAITS)
        + "\n\n"
        "STEP 2 & 3 — Variable Life Path + The Math:\n"
        "Generate the genome_mutation_log. This is an array — it may be empty [], "
        "have 1 event, or many events. Do NOT force one event per stage.\n"
        "For each event that DID happen, provide specific +/- integer modifiers.\n"
        "Examples of valid entries:\n"
        '  {"life_stage": "Formation", "event_description": "Father\'s business failed; dropped out at 16", '
        '"trait_modifiers": {"conscientiousness": -12, "socioeconomic_friction": 20, "literacy_and_articulation": -15}}\n'
        '  {"life_stage": "Independence", "event_description": "Got motorbike loan; defaulted after 8 months", '
        '"trait_modifiers": {"neuroticism": 10, "brand_relationship": -8, "socioeconomic_friction": 18}}\n'
        "A completely stable childhood with no major events = empty Origin stage (exclude it from the log).\n"
        "Let chaos be chaotic and stability be uneventful."
    )
    mutation_result = _genesis_call_tool(
        anthropic_client, _MUTATION_TOOL, mutation_msg,
        system_prompt=system_prompt, max_tokens=1200,
    )
    if mutation_result is None:
        print("    Genesis: mutation log failed — using base genome as-is")
        mutation_log = []
    else:
        mutation_log = mutation_result.get("genome_mutation_log") or []

    # ── STEP 4 (Python math): apply modifiers → Final Genome ────────────────
    final_genome = _apply_mutations(base_genome, mutation_log)
    print(
        f"    Mutations applied: {len(mutation_log)} event(s) across "
        f"{len({e.get('life_stage') for e in mutation_log})} stage(s)"
    )

    # ── CALL 3: Life blueprint ─────────────────────────────────────────
    blueprint_msg = (
        f"CITY FOR THIS AGENT: {city}\n"
        f"Agent profile: {seed}\n\n"
        f"Final (mutated) Personality: "
        + ", ".join(f"{k}={final_genome.get(k)}" for k in _MUTABLE_TRAITS)
        + f"\nKey life events that caused these mutations:\n"
        + "\n".join(f"  [{e['life_stage']}] {e['event_description']}" for e in mutation_log)
        + f"\n\nGenerate ALL 5 life layers rooted in {city}. STRICT BREVITY: summary ≤15 words, "
        "key_events 3 items each max 8 words, psychological_impact max 12 words. "
        "All 5 layers must fit in one response."
    )
    blueprint = _genesis_call_tool(
        anthropic_client, _BLUEPRINT_TOOL, blueprint_msg,
        system_prompt=system_prompt, max_tokens=1500,
    )
    if blueprint is None:
        return None
    for layer in ["origin_layer", "formation_layer", "independence_layer", "maturity_layer", "legacy_layer"]:
        if not isinstance(blueprint.get(layer), dict):
            return None

    # Merge: final_genome integers + voice_print + blueprint layers + mutation log
    result = dict(final_genome)
    result.update(blueprint)
    result["genome_mutation_log"] = mutation_log   # stored as JSONB
    return result


def dynamic_seed_agents(
    demographic: str,
    count: int,
    sb,                  # supabase client
    anthropic_client,    # anthropic client
    city_pool: list[str] | None = None,  # override for non-national demographics
    seed_templates: list[str] | None = None, # override templates
    country: str = "Indonesia",          # override country name
) -> list[dict]:
    """
    Generate `count` new agents for `demographic`, insert them into Supabase,
    and return the list of newly created agent row dicts.

    Geographic diversity is enforced via two mechanisms:
      1. Python randomly picks a city from `city_pool` (or _INDONESIAN_CITIES) BEFORE
         any LLM call — the model cannot override this choice.
      2. The city is injected into both the system prompt and the user message,
         with an explicit CRITICAL directive not to default to any other location.
    """
    # Parse age range from demographic string, e.g. '25-45' → (25, 45)
    age_match = re.search(r"(\d+)[^\d]+(\d+)", demographic)
    age_min, age_max = (int(age_match.group(1)), int(age_match.group(2))) if age_match else (25, 45)

    pool = city_pool if city_pool else _INDONESIAN_CITIES
    templates = seed_templates if seed_templates else _SEED_TEMPLATES

    print(f"  Genesis: spawning {count} new agent(s) for '{demographic}'...")
    created: list[dict] = []
    SLEEP_BETWEEN = 70  # seconds — stay inside 10K output TPM window

    for i in range(count):
        # ── Python picks city deterministically — LLM cannot override ─────────
        city = random.choice(pool)
        age = random.randint(age_min, age_max)
        template = random.choice(templates)
        seed = template.format(age=age, city=city)
        region = f"{city}, {country}" if city != country else city

        print(f"  Genesis [{i + 1}/{count}] Age {age} | City: {city}")
        if i > 0:
            print(f"    Sleeping {SLEEP_BETWEEN}s (rate-limit buffer)...")
            time.sleep(SLEEP_BETWEEN)

        data = _generate_one_agent(anthropic_client, age, seed, city=city, demographic=demographic)
        if data is None:
            print(f"    Genesis agent {i + 1} FAILED — skipping")
            continue

        payload = {
            "id":                       str(uuid.uuid4()),
            "created_at":               datetime.now(timezone.utc).isoformat(),
            "target_demographic":       demographic,
            "age":                      age,
            "region":                   region,
            "openness":                 data["openness"],
            "conscientiousness":        data["conscientiousness"],
            "extraversion":             data["extraversion"],
            "agreeableness":            data["agreeableness"],
            "neuroticism":              data["neuroticism"],
            "communication_style":      data["communication_style"],
            "decision_making":          data["decision_making"],
            "brand_relationship":       data["brand_relationship"],
            "influence_susceptibility": data["influence_susceptibility"],
            "emotional_expression":     data["emotional_expression"],
            "conflict_behavior":        data["conflict_behavior"],
            "literacy_and_articulation": data.get("literacy_and_articulation", 50),
            "socioeconomic_friction":   data.get("socioeconomic_friction", 50),
            "identity_fusion":          data.get("identity_fusion", 50),
            "chronesthesia_capacity":   data.get("chronesthesia_capacity", 50),
            "tom_self_awareness":       data.get("tom_self_awareness", 50),
            "tom_social_modeling":      data.get("tom_social_modeling", 50),
            "executive_flexibility":    data.get("executive_flexibility", 50),
            "religion":                 data.get("religion", "Unspecified"),
            "cultural_background":      data.get("cultural_background", "Unspecified"),
            "genome_mutation_log":      data.get("genome_mutation_log", []),  # JSONB array
            "origin_layer":             data["origin_layer"],
            "formation_layer":          data["formation_layer"],
            "independence_layer":       data["independence_layer"],
            "maturity_layer":           data["maturity_layer"],
            "legacy_layer":             data["legacy_layer"],
            "voice_print":              data["voice_print"],
        }
        sb.table("agent_genomes").insert(payload).execute()
        print(f"    Genesis agent {i + 1} stored (Age {age}, all columns filled)")
        created.append(payload)

    return created


# ─────────────────────────────────────────────────────────────────────────────
# SEMANTIC ROUTER  — Claude Haiku evaluates which existing demographics are
# 'within reason' of the requested PTM/STM before counting or seeding agents.
# ─────────────────────────────────────────────────────────────────────────────

_SEMANTIC_ROUTER_TOOL = {
    "name": "return_matching_demographics",
    "description": (
        "Return ONLY the demographic strings from the provided list that are "
        "practically synonymous with or a reasonable substitute for the requested PTM/STM. "
        "Do NOT include demographics that are clearly different in age, class, or geography."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "approved_matches": {
                "type": "array",
                "description": (
                    "Array of approved demographic strings, chosen verbatim from the "
                    "available_demographics list. May be empty if nothing matches."
                ),
                "items": {"type": "string"}
            }
        },
        "required": ["approved_matches"],
        "additionalProperties": False,
    },
}


def evaluate_demographics(
    ptm_requested: str,
    stm_requested: str,
    available_demos: list[str],
    anthropic_client,
) -> list[str]:
    """
    Use Claude Haiku as a Semantic Router to identify which existing demographics
    in our database are 'within reason' of the requested PTM / STM markets.

    Returns a list of approved demographic strings (subset of available_demos).
    Falls back to returning the requested demographics themselves if the call fails,
    so the deficit logic can still seed fresh agents correctly.
    """
    import anthropic as _anthropic

    if not available_demos:
        print("  Semantic Router: no existing demographics in DB — skipping evaluation.")
        return []

    stm_line = f"STM: {stm_requested}" if stm_requested.strip() else "STM: (none)"
    available_str = "\n".join(f"  - {d}" for d in available_demos)

    system_prompt = (
        "You are a semantic routing agent for a focus group research database.\n"
        "Your job is to find existing agent pools that are CLOSE ENOUGH to the "
        "requested target markets to be used without generating new agents from scratch.\n\n"
        "Matching criteria (any of these justify an approve):\n"
        "  - Same broad age range (±5 years is fine, ±15 years is NOT)\n"
        "  - Same socioeconomic class (Upper, Upper-Middle, Middle, Working — must match)\n"
        "  - Same country or major region (Indonesian city-to-city is OK; cross-country is NOT)\n"
        "  - Similar life stage or digital behaviour (e.g., 'Urban Millennial' ≈ 'Young Professional')\n\n"
        "When in doubt, be INCLUSIVE — it is better to reuse a close pool than spawn duplicates.\n"
        "Never include demographics that clearly differ in class, age bracket, or country."
    )

    user_message = (
        f"The user needs focus group agents for these markets:\n"
        f"  PTM: {ptm_requested}\n"
        f"  {stm_line}\n\n"
        f"Here are ALL the demographics currently available in our database:\n"
        f"{available_str}\n\n"
        f"Return ONLY the demographic strings from that list that are practically "
        f"synonymous with or a reasonable substitute for the PTM or STM above."
    )

    try:
        response = anthropic_client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=512,
            system=system_prompt,
            tools=[_SEMANTIC_ROUTER_TOOL],
            tool_choice={"type": "tool", "name": "return_matching_demographics"},
            messages=[{"role": "user", "content": user_message}],
        )
        block = next((b for b in response.content if b.type == "tool_use"), None)
        if block is None:
            raise ValueError("No tool_use block returned by semantic router")
        approved = block.input.get("approved_matches") or []
        # Safety: only keep strings that actually exist in the available list
        valid = [d for d in approved if d in available_demos]
        print(
            f"  Semantic Router ✓ — {len(valid)} match(es) approved from "
            f"{len(available_demos)} available: {valid or '(none)'}"
        )
        return valid
    except Exception as e:
        print(f"  Semantic Router ERROR: {e} — proceeding with empty match list.")
        return []



# ────────────────────────────────────────────────────────────────────────────
# NODE 1: Intake, Semantic Routing, Check-and-Fill, Query
#         (local — runs on the caller's machine)
# ────────────────────────────────────────────────────────────────────────────

def node1_intake_and_query(target_demographic: str, limit: int = 10) -> list[dict]:
    """
    Semantic-routing intake:
      1. Fetch all unique target_demographic strings from Supabase.
      2. Run the Semantic Router (Claude Haiku) to find existing demographics
         that are 'within reason' of the requested PTM / STM.
      3. Count agents belonging to the approved demographics.
      4. If deficit > 0, generate the missing agents using the *original*
         requested demographics (so exact matches accumulate over time).
      5. Fetch a large pool from approved + newly seeded demographics,
         random-sample `limit` agents, and return.

    target_demographic may be pipe-separated for PTM|STM:
      "Jakarta Young Professional, 22-28|Bandung Middle Class, 30-40"
    """
    from dotenv import load_dotenv
    from supabase import create_client
    import anthropic as _anthropic

    load_dotenv("pantheon.env", override=True)
    url = os.environ.get("SUPABASE_URL")
    key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY")
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not url or not key:
        raise RuntimeError("Missing SUPABASE_URL or SUPABASE_SERVICE_ROLE_KEY")

    sb = create_client(url, key)
    anthropic_client = _anthropic.Anthropic(api_key=api_key)

    # Parse PTM / STM from pipe-separated string
    parts = [d.strip() for d in target_demographic.split("|") if d.strip()]
    ptm = parts[0] if len(parts) >= 1 else target_demographic.strip()
    stm = parts[1] if len(parts) >= 2 else ""
    requested_demos = [d for d in parts if d]  # all requested (PTM + optional STM)
    label = " + ".join(requested_demos)

    print(f"Node 1: Semantic Routing — PTM='{ptm}' | STM='{stm or 'none'}'")

    # ── Step 1: Fetch all unique demographics in the database ─────────────────
    unique_q = sb.table("agent_genomes").select("target_demographic")
    unique_result = unique_q.execute()
    all_demos_raw = [row["target_demographic"] for row in (unique_result.data or []) if row.get("target_demographic")]
    available_demos = sorted(set(all_demos_raw))
    print(f"  DB contains {len(available_demos)} unique demographic(s): {available_demos}")

    # ── Step 2: Semantic Router — find close-enough existing demographics ──────
    approved_demos = evaluate_demographics(ptm, stm, available_demos, anthropic_client)

    # Always include the exact requested demographics in the approved set
    # (they may have been just seeded or already exist with exact string match).
    for d in requested_demos:
        if d not in approved_demos:
            approved_demos.append(d)

    # ── Step 3: Count agents in the approved pool ─────────────────────────────
    if approved_demos:
        count_q = sb.table("agent_genomes").select("id", count="exact")
        count_q = count_q.in_("target_demographic", approved_demos)
        count_result = count_q.execute()
        existing_count = count_result.count or 0
    else:
        existing_count = 0

    print(f"Node 1: Approved pool has {existing_count} agent(s) across {len(approved_demos)} demographic(s) (need {limit}).")

    # ── Step 4: Deficit fill using the original requested demographics ────────
    deficit = limit - existing_count
    if deficit > 0:
        print(f"Node 1: Deficit = {deficit} agent(s). Triggering Genesis Protocol...")
        # Spread generation evenly across all requested (PTM + STM)
        per_demo = max(1, deficit // max(len(requested_demos), 1))
        remainder = deficit - (per_demo * len(requested_demos))
        for idx, demo in enumerate(requested_demos):
            n = per_demo + (1 if idx < remainder else 0)
            if n > 0:
                newly_created = dynamic_seed_agents(demo, n, sb, anthropic_client)
                # Add the newly seeded demo to approved so it's fetched below
                if demo not in approved_demos:
                    approved_demos.append(demo)
        print(f"Node 1: Genesis complete. {deficit} new agent(s) created.")
    else:
        print(f"Node 1: Pool sufficient — no genesis needed.")

    # ── Step 5: Fetch full pool and random-sample `limit` agents ─────────────
    fetch_q = sb.table("agent_genomes").select("*")
    if approved_demos:
        fetch_q = fetch_q.in_("target_demographic", approved_demos)

    pool_result = fetch_q.limit(500).execute()
    pool = pool_result.data or []

    selected = random.sample(pool, min(limit, len(pool)))

    print(f"Node 1 ✓ — returning {len(selected)} agent(s) for pipeline (pool size: {len(pool)}).")

    # ── Emit agent sample for dashboard Agent Inspector ───────────────────────
    _INSPECTOR_FIELDS = [
        "target_demographic", "age", "region",
        "openness", "conscientiousness", "extraversion", "agreeableness", "neuroticism",
        "communication_style", "decision_making", "brand_relationship",
        "influence_susceptibility", "emotional_expression", "conflict_behavior",
        "literacy_and_articulation", "socioeconomic_friction",
        "identity_fusion", "chronesthesia_capacity", "tom_self_awareness",
        "tom_social_modeling", "executive_flexibility",
        "genome_mutation_log", "voice_print",
    ]
    sample_agents = selected
    sample_slim = [{k: a.get(k) for k in _INSPECTOR_FIELDS} for a in sample_agents]
    print(f"PANTHEON_AGENT_SAMPLE::{json.dumps(sample_slim, ensure_ascii=False)}")

    return selected



# ─────────────────────────────────────────────────────────────────────────────
# NODE 2: Runtime Snapshot  (parallel Haiku via .map())
# ─────────────────────────────────────────────────────────────────────────────

NODE2_TOOL = {
    "name": "submit_snapshot",
    "description": "Submit the agent's current runtime mental/emotional/financial state.",
    "input_schema": {
        "type": "object",
        "properties": {
            "current_emotional_state": {
                "type": "string",
                "description": "1 sentence: specific emotional state right now, before a focus group."
            },
            "current_mental_bandwidth": {
                "type": "string",
                "description": "One of: Exhausted / Scattered / Focused / Highly focused / Distracted / Anxious"
            },
            "current_financial_pressure": {
                "type": "string",
                "description": "1 sentence: immediate financial context or pressure they're experiencing."
            }
        },
        "required": ["current_emotional_state", "current_mental_bandwidth", "current_financial_pressure"],
        "additionalProperties": False
    }
}

NODE2_SYSTEM = (
    "You are PANTHEON's Runtime State Engine. Given an agent's life history and personality genome, "
    "generate their precise mental and emotional state RIGHT NOW — as they are about to participate "
    "in a consumer focus group session. Be psychologically specific. Reference their actual life stage "
    "events, financial trajectory, and neuroticism/conscientiousness levels. Not generic.\n\n"
    "CHRONESTHESIA: Agents with high Chrono scores (70+) show anticipatory thinking about how this "
    "focus group topic fits their long-term life plans — they may arrive with pre-formed strategic "
    "questions. Low Chrono agents (<30) focus on immediate state: tiredness, hunger, mood.\n"
    "EXECUTIVE FLEXIBILITY: High ExecFlex agents may mask their true emotional state behind a "
    "professional facade. Low ExecFlex agents' runtime state leaks visibly."
)


@app.function(image=image, secrets=SECRETS, timeout=120)
def node2_generate_snapshot(agent: dict) -> dict:
    """Generate the 'right now' runtime mental state for one agent via Claude Haiku."""
    import anthropic as _anthropic

    client = _anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    ctx = _build_agent_context(agent)

    prompt = (
        f"{ctx}\n\n"
        "Based on this person's history and personality, describe their exact mental and emotional "
        "state RIGHT NOW, just before a consumer focus group session begins."
    )

    try:
        resp = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=300,
            system=NODE2_SYSTEM,
            tools=[NODE2_TOOL],
            tool_choice={"type": "tool", "name": "submit_snapshot"},
            messages=[{"role": "user", "content": prompt}]
        )
        block = next((b for b in resp.content if b.type == "tool_use"), None)
        snapshot = block.input if block else {
            "current_emotional_state": "Unknown",
            "current_mental_bandwidth": "Unknown",
            "current_financial_pressure": "Unknown"
        }
    except Exception as e:
        print(f"  Node 2 snapshot error for agent {agent.get('id', '')[:8]}: {e}")
        snapshot = {
            "current_emotional_state": "Unable to generate",
            "current_mental_bandwidth": "Unknown",
            "current_financial_pressure": "Unknown"
        }

    print(f"  Node 2 ✓ | Age {agent.get('age')} | {snapshot['current_mental_bandwidth']}")
    return {"immutable_dna": agent, "runtime_snapshot": snapshot}


# ─────────────────────────────────────────────────────────────────────────────
# NODE 3: Phase A — Mass Session  (parallel Haiku via .starmap())
# ─────────────────────────────────────────────────────────────────────────────

NODE3_TOOL = {
    "name": "submit_phase_a_reaction",
    "description": "Submit this agent's immediate gut reaction to the campaign stimulus.",
    "input_schema": {
        "type": "object",
        "properties": {
            "gut_reaction": {
                "type": "string",
                "description": "1 sentence: visceral, psychologically authentic first reaction."
            },
            "emotional_temperature": {
                "type": "integer",
                "description": "1 (cold/unresponsive) to 10 (highly emotionally activated)."
            },
            "dominant_emotion": {
                "type": "string",
                "description": "Single word or short phrase."
            },
            "personal_relevance_score": {
                "type": "integer",
                "description": "1 (irrelevant) to 10 (directly speaks to my situation)."
            },
            "intent_signal": {
                "type": "string",
                "description": "Exactly one of: ignore, glance, click, save, share, complain, dismiss."
            }
        },
        "required": [
            "gut_reaction", "emotional_temperature", "dominant_emotion",
            "personal_relevance_score", "intent_signal"
        ],
        "additionalProperties": False
    }
}

NODE3_SYSTEM = (
    "You are PANTHEON Phase A Engine. Simulate the exact immediate reaction of a specific "
    "consumer when they encounter an ad on their phone.\n"
    "Stay fully in character. The reaction must be consistent with: "
    "their personality integers, life history, runtime emotional state, financial situation, "
    "and cultural background.\n"
    "Do NOT generate a generic consumer response. This is a specific human with a specific history.\n\n"
    "IDENTITY FUSION: High IdFusion agents (70+) filter gut reactions through their group identity — "
    "their first thought is 'What would my family/community think?' or 'Does this align with our values?' "
    "Low IdFusion agents (<30) react personally: 'Do I need this? Is this for me?'\n"
    "CHRONESTHESIA: High Chrono agents (70+) immediately project forward: 'Where does this fit in my "
    "5-year plan?' Low Chrono agents react to the immediate stimulus without future reference."
)


@app.function(image=image, secrets=SECRETS, timeout=120)
def node3_mass_session(dynamic_agent: dict, campaign_brief: str) -> dict:
    """Phase A — one agent's uninfluenced first reaction to the campaign stimulus."""
    import anthropic as _anthropic

    client = _anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    dna = dynamic_agent["immutable_dna"]
    snap = dynamic_agent["runtime_snapshot"]
    ctx = _build_agent_context(dna)

    prompt = (
        f"{ctx}\n\n"
        f"Runtime state:\n"
        f"  Emotional state: {snap.get('current_emotional_state')}\n"
        f"  Mental bandwidth: {snap.get('current_mental_bandwidth')}\n"
        f"  Financial pressure: {snap.get('current_financial_pressure')}\n\n"
        f"Campaign stimulus: {campaign_brief}\n\n"
        "Simulate this agent's immediate gut reaction to this ad."
    )

    reaction = None
    status = "error"
    error = None
    for _attempt in range(1, 4):
        try:
            resp = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=300,
                system=NODE3_SYSTEM,
                tools=[NODE3_TOOL],
                tool_choice={"type": "tool", "name": "submit_phase_a_reaction"},
                messages=[{"role": "user", "content": prompt}]
            )
            block = next((b for b in resp.content if b.type == "tool_use"), None)
            if not block:
                raise ValueError("No tool_use block returned")
            reaction = block.input
            status = "ok"
            error = None
            break
        except _anthropic.RateLimitError as e:
            wait = 65
            try:
                wait = int(e.response.headers.get("retry-after", 65)) + 2
            except Exception:
                pass
            print(f"  Node 3: rate limit hit (attempt {_attempt}/3) — sleeping {wait}s...")
            time.sleep(wait)
            error = str(e)
        except Exception as e:
            error = str(e)
            print(f"  Node 3 error for agent {dna.get('id', '')[:8]} (attempt {_attempt}/3): {e}")
            if _attempt < 3:
                time.sleep(5)
    if status == "error":
        print(f"  Node 3 ✗ agent {dna.get('id', '')[:8]}: all retries exhausted — {error}")

    print(
        f"  Node 3 ✓ | Age {dna.get('age')} | "
        f"{reaction.get('dominant_emotion', 'ERROR') if reaction else 'ERROR'} | "
        f"Intent: {reaction.get('intent_signal', '-') if reaction else '-'}"
    )
    return {
        "agent_id": dna.get("id"),
        "age": dna.get("age"),
        "demographic": dna.get("target_demographic"),
        "agent_context": ctx,
        "runtime_snapshot": snap,
        "phase_a": reaction,
        "status": status,
        "error": error,
    }


# ─────────────────────────────────────────────────────────────────────────────
# NODE 4: Phase B — Breakout Rooms  (parallel Sonnet via .starmap())
# ─────────────────────────────────────────────────────────────────────────────

NODE4_SYSTEM = (
    "You are a qualitative research director facilitating a focus group simulation for a "
    "consumer research firm.\n\n"
    "Simulate how THESE SPECIFIC PEOPLE actually talk — not generic consumers. "
    "Use their personality scores, cultural backgrounds, socioeconomic realities, and "
    "literacy_and_articulation scores to shape HOW they speak.\n\n"
    "FORMAT RULES (mandatory for every speaking turn):\n"
    "[Agent Name/Age] (Spoken): What they actually say aloud. "
    "If literacy_and_articulation is below 35, they must use simple, broken, or halting language. "
    "If they are intimidated or deferential (high agreeableness, low extraversion), "
    "they speak briefly or trail off.\n"
    "[Agent Name/Age] (Inner Thought): What they actually think but withhold. "
    "This MUST reveal whether they are silently conforming, secretly unconvinced, "
    "suppressing anger, or feeling shame. This is the layer the spoken word hides.\n\n"
    "THEORY OF MIND DYNAMICS (use ToMSelf and ToMSocial scores):\n"
    "- High ToMSelf + Low ToMSocial: authentic but socially clumsy — says what they think even "
    "when inappropriate, genuinely unaware of how it lands.\n"
    "- Low ToMSelf + High ToMSocial: performative — adjusts speech to match what they think others "
    "want to hear, but inner thoughts reveal they don't actually know what they themselves believe.\n"
    "- High both: socially sophisticated — spoken words are calibrated, inner thoughts are strategic.\n"
    "- Low both: impulsive, reactive — spoken words and inner thoughts are nearly identical.\n\n"
    "EXECUTIVE FLEXIBILITY (use ExecFlex score to control the GAP between inner thought and spoken word):\n"
    "- High ExecFlex (70+): large gap possible — the professional who is privately seething but speaks "
    "calmly. Can suppress base trait impulses (neuroticism, conflict_behavior) in social settings.\n"
    "- Low ExecFlex (<30): small gap — what they think leaks directly into what they say. "
    "High N + Low ExecFlex = visibly anxious, voice trembles, cannot hide distress.\n"
    "- High N + High ExecFlex = appears calm but inner thoughts are catastrophizing.\n\n"
    "IDENTITY FUSION (use IdFusion score for group loyalty behavior):\n"
    "- High IdFusion (70+): will defend in-group positions even against personal interest. "
    "When another participant from their cultural/religious group is challenged, high-fusion "
    "agents rally to defend REGARDLESS of their own opinion on the topic.\n"
    "- Low IdFusion (<30): evaluates arguments purely on personal merit, not group loyalty.\n\n"
    "Do not force consensus or resolution. Let class anxiety, financial shame, and "
    "social pressure drive the subtext. Make it feel like raw, real fieldwork audio."
)


@app.function(image=image, secrets=SECRETS, timeout=600)
def node4_breakout_room(group_reactions: list[dict], campaign_brief: str) -> dict:
    """Phase B — simulate one breakout room debate transcript for a group of agents."""
    import anthropic as _anthropic

    client = _anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    # Build group context
    parts = []
    for r in group_reactions:
        pa = r.get("phase_a") or {}
        snap = r.get("runtime_snapshot") or {}
        parts.append(
            f"PARTICIPANT (Age {r.get('age')}, {r.get('demographic')})\n"
            f"Profile: {r.get('agent_context', 'N/A')}\n"
            f"Pre-session state: {snap.get('current_emotional_state', 'N/A')}\n"
            f"Initial reaction: \"{pa.get('gut_reaction', 'N/A')}\"\n"
            f"Emotion: {pa.get('dominant_emotion', 'N/A')} | "
            f"Temp: {pa.get('emotional_temperature', '-')}/10 | "
            f"Relevance: {pa.get('personal_relevance_score', '-')}/10 | "
            f"Intent: {pa.get('intent_signal', '-')}"
        )

    group_ctx = "\n\n".join(parts)
    agent_ages = [str(r.get("age")) for r in group_reactions]

    prompt = (
        f"CAMPAIGN BRIEF:\n{campaign_brief}\n\n"
        f"{'─' * 50}\n"
        f"FOCUS GROUP PARTICIPANTS:\n\n{group_ctx}\n\n"
        f"{'─' * 50}\n"
        f"INSTRUCTIONS:\n"
        f"These {len(group_reactions)} people (ages: {', '.join(agent_ages)}) have just seen the "
        f"ad above and shared their initial reactions. Now they are in a focus group breakout room.\n\n"
        "Write a raw, realistic 15-turn debate transcript. "
        "MANDATORY DUAL-CHANNEL FORMAT for every single turn:\n\n"
        "[Agent Name/Age] (Spoken): What they actually say aloud — constrained STRICTLY by their "
        "literacy_and_articulation score. Score 1-35: simple, broken or halting speech, short sentences, "
        "filler words, may trail off. Score 36-60: functional but plain everyday language. "
        "Score 61-80: articulate, expressive. Score 81-100: eloquent, complex sentences, precise vocabulary.\n"
        "[Agent Name/Age] (Inner Thought): What they actually think but withhold due to social pressure, "
        "lack of confidence, agreeableness, class anxiety, or financial shame. "
        "MUST reveal if they are silently conforming, secretly unconvinced, suppressing anger, "
        "feeling envy, or hiding ignorance about the product. "
        "This is the layer the spoken word conceals.\n\n"
        "Additional rules:\n"
        "- Let them interrupt, disagree, go on tangents, reveal personal details under social pressure\n"
        "- They must reference their specific financial realities, debts, and socioeconomic friction\n"
        "- Higher socioeconomic_friction agents show more suspicion, fatalism, or desperation\n"
        "- Do NOT force consensus — real focus groups are messy\n"
        "- Make the inner thoughts the most honest and disturbing part of the transcript"
    )

    try:
        resp = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=2500,
            system=NODE4_SYSTEM,
            messages=[{"role": "user", "content": prompt}]
        )
        transcript = resp.content[0].text
        status = "ok"
        error = None
    except Exception as e:
        transcript = None
        status = "error"
        error = str(e)
        print(f"  Node 4 error: {e}")

    print(f"  Node 4 ✓ | Group ({', '.join(agent_ages)}) | status={status}")
    return {
        "participant_ages": agent_ages,
        "transcript": transcript,
        "status": status,
        "error": error,
    }


# ─────────────────────────────────────────────────────────────────────────────
# NODE 5: Phase C — Synthesis  (single Sonnet call)
# ─────────────────────────────────────────────────────────────────────────────

NODE5_SYSTEM = (
    "You are PANTHEON — the Team Lead and chief synthesis intelligence for a qualitative "
    "research firm specialising in Indonesian consumer psychology.\n\n"
    "You receive: (1) Phase A mass reaction data from multiple agents, "
    "(2) Phase B debate transcripts from focus group breakout rooms.\n\n"
    "Your output is the final Research Intelligence Report. "
    "Be incisive, specific, and actionable. No filler. No hedging. "
    "Write like a brilliant strategist, not a consultant. "
    "Reference specific agents, direct quotes, and data points from the input."
)

REPORT_PROMPT = """\
Generate the PANTHEON Research Intelligence Report with EXACTLY these 7 sections:

1. THE HEADLINE TRUTH
   One brutal, specific sentence capturing the single most important consumer truth revealed.
   Not a finding — a truth.

2. PTM SIGNAL ANALYSIS (Primary Target Market)
   Who actually responded positively? Define the PTM precisely from the data.
   What specific emotional and financial conditions make someone receptive?
   Quote and cite specific agents with their age and demographic.

3. STM SIGNAL ANALYSIS (Secondary Target Market)
   Who showed unexpected or conditional interest?
   What specific trigger would convert them? Be precise.

4. THE FRACTURE LINES
   What were the deepest points of disagreement in the focus group debates?
   Identify 2-3 specific fault lines. What does each side's position reveal about their psychology?

5. THE INVISIBLE INSIGHT
   What did the agents NOT say directly but reveal through their behavior, word choice,
   and emotional temperature? This is the insight the brand would miss reading only surface data.

6. THREE SHARPENING RECOMMENDATIONS
   Three specific, actionable campaign changes. Each must reference specific agent reactions as evidence.
   Format: [RECOMMENDATION] → [EVIDENCE] → [EXPECTED RESPONSE SHIFT]

7. THE KILL SWITCH
   What is the single biggest risk that could make this campaign backfire catastrophically?
   Name it precisely. Who is most at risk of becoming an active detractor and why?\
"""


@app.function(image=image, secrets=SECRETS, timeout=1200)   # 20-min window
def node5_synthesis(
    mass_reactions: list[dict],
    breakout_transcripts: list[dict],
    campaign_brief: str,
) -> str:
    """Phase C — synthesise all data into the final 7-section PANTHEON Report."""
    import anthropic as _anthropic

    client = _anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])
    print("Node 5: Building synthesis payload...")

    # Phase A summary
    phase_a_lines = []
    for r in mass_reactions:
        if r.get("status") != "ok" or not r.get("phase_a"):
            continue
        pa = r["phase_a"]
        phase_a_lines.append(
            f"Agent Age {r.get('age')} ({r.get('demographic')}): "
            f"Emotion={pa.get('dominant_emotion')} | Temp={pa.get('emotional_temperature')}/10 | "
            f"Relevance={pa.get('personal_relevance_score')}/10 | Intent={pa.get('intent_signal')}\n"
            f"  Gut reaction: \"{pa.get('gut_reaction')}\""
        )

    # Phase B transcripts
    phase_b_lines = []
    for tb in breakout_transcripts:
        if tb.get("status") != "ok" or not tb.get("transcript"):
            phase_b_lines.append(f"[Group transcript unavailable — {tb.get('error')}]")
        else:
            ages = ", ".join(tb.get("participant_ages", []))
            phase_b_lines.append(
                f"GROUP (Ages: {ages}):\n{tb['transcript']}"
            )

    payload = (
        f"CAMPAIGN BRIEF:\n{campaign_brief}\n\n"
        f"{DIVIDER}\n"
        f"PHASE A — MASS SESSION RESULTS ({len(phase_a_lines)} agents):\n\n"
        + "\n\n".join(phase_a_lines)
        + f"\n\n{DIVIDER}\n"
        f"PHASE B — BREAKOUT ROOM TRANSCRIPTS:\n\n"
        + f"\n\n{'─'*50}\n\n".join(phase_b_lines)
        + f"\n\n{DIVIDER}\n\n"
        + REPORT_PROMPT
    )

    print("Node 5: Firing synthesis call (claude-sonnet-4-5, max_tokens=8192)...")

    # ── Conversation history for the auto-resume loop ─────────────────────────
    messages = [{"role": "user", "content": payload}]
    final_report = ""
    chunk_num = 0

    while True:
        chunk_num += 1
        resp = client.messages.create(
            model="claude-sonnet-4-5",
            max_tokens=8192,
            system=NODE5_SYSTEM,
            messages=messages,
            extra_headers={"anthropic-beta": "max-tokens-3-5-sonnet-2024-07-15"},
        )

        chunk_text = resp.content[0].text if resp.content else ""
        final_report += chunk_text
        stop_reason = resp.stop_reason

        print(
            f"Node 5: chunk {chunk_num} — {len(chunk_text):,} chars | "
            f"stop_reason={stop_reason} | total={len(final_report):,} chars"
        )

        if stop_reason != "max_tokens":
            # 'end_turn' or unexpected — report is complete
            break

        # ── Auto-resume: model hit token cap mid-response ─────────────────────
        print("Node 5: max_tokens hit — firing continuation call...")
        # Append assistant's partial chunk to history, then ask it to continue.
        messages.append({"role": "assistant", "content": chunk_text})
        messages.append({
            "role": "user",
            "content": (
                "You hit your length limit. Please continue writing the report exactly "
                "from where you left off. Do not apologize or add transition text, "
                "just continue the next word."
            ),
        })

    return final_report


# ─────────────────────────────────────────────────────────────────────────────
# ─────────────────────────────────────────────────────────────────────────────
# PIPELINE EXECUTION CORE
# ─────────────────────────────────────────────────────────────────────────────

def run_pipeline_core(
    target: str = "Medanese Upper Middle Class, 25-45",
    brief: str = (
        "A fintech app displays an iPhone advertisement for a Buy-Now-Pay-Later fintech app "
        "that lets users split rent payments into 4 installments in Indonesia."
    ),
    brief_file: str = "",
    client: str = "",
    limit: int = 10,
    group_size: int = 5,
):
    """
    Run the full PANTHEON 5-node pipeline on Modal.

    Args:
      target     — target_demographic to query from Supabase
      brief      — campaign brief text
      brief_file — path to a file containing the brief text (to avoid cmd length limits)
      client     — optional client name for report output filename
      limit      — number of agents to pull (default 10)
      group_size — agents per Phase B breakout room (default 5)
    """
    if brief_file:
        try:
            with open(brief_file, "r", encoding="utf-8") as f:
                brief = f.read().strip()
        except Exception as e:
            print(f"Failed to load brief from {brief_file}: {e}")
            import sys
            sys.exit(1)

    print(f"\n{DIVIDER}")
    print("  PANTHEON — Full Pipeline (Modal Serverless)")
    print(f"{DIVIDER}")
    print(f"  Target      : {target}")
    print(f"  Limit       : {limit} agents")
    print(f"  Group size  : {group_size} per breakout room")
    print(f"  Brief       : {brief[:80]}...")
    print(f"{DIVIDER}\n")

    # ── Node 1: Query ──────────────────────────────────────────────────────
    agents = node1_intake_and_query(target, limit)
    if not agents:
        return {"error": "No agents found. Run seed_genomes.py first."}

    # ── Node 2: Runtime snapshots via .map() ───────────────────────────────
    print(f"\nNode 2 — Generating runtime snapshots ({len(agents)} agents in parallel)...")
    dynamic_agents = list(node2_generate_snapshot.map(agents))
    print(f"  ✓ {len(dynamic_agents)} snapshots complete.\n")

    # ── Node 3: Phase A via .starmap() ─────────────────────────────────────
    print(f"Node 3 — Phase A mass session ({len(dynamic_agents)} agents in parallel)...")
    node3_args: list[tuple[Any, ...]] = [(da, brief) for da in dynamic_agents]
    mass_reactions = list(node3_mass_session.starmap(node3_args))
    ok_a = sum(1 for r in mass_reactions if r["status"] == "ok")
    print(f"  ✓ {ok_a}/{len(mass_reactions)} Phase A reactions collected.\n")

    for r in mass_reactions:
        pa = r.get("phase_a") or {}
        mark = "✓" if r["status"] == "ok" else "✗"
        print(
            f"  {mark} Age {r.get('age'):>2} | {pa.get('dominant_emotion', 'ERROR'):<35} | "
            f"Temp {pa.get('emotional_temperature', '-')}/10 | Intent: {pa.get('intent_signal', '-')}"
        )

    # ── Node 4: Phase B breakout rooms via .starmap() ──────────────────────
    groups: list[list[dict]] = [
        mass_reactions[i : i + group_size]
        for i in range(0, len(mass_reactions), group_size)
    ]
    print(f"\nNode 4 — Phase B breakout rooms ({len(groups)} groups, parallel Sonnet calls)...")
    node4_args: list[tuple[Any, ...]] = [(g, brief) for g in groups]
    breakout_transcripts = list(node4_breakout_room.starmap(node4_args))
    ok_b = sum(1 for t in breakout_transcripts if t["status"] == "ok")
    print(f"  ✓ {ok_b}/{len(breakout_transcripts)} transcripts generated.\n")

    for i, tb in enumerate(breakout_transcripts, 1):
        print(f"{DIVIDER}")
        print(f"  PHASE B — GROUP {i} TRANSCRIPT (Ages: {', '.join(tb['participant_ages'])})")
        print(f"{DIVIDER}")
        if tb["status"] == "ok":
            print(tb["transcript"])
        else:
            print(f"  ERROR: {tb.get('error')}")
        print()

    # ── Node 5: Phase C synthesis ───────────────────────────────────────────
    print(f"{DIVIDER}")
    print("Node 5 — Phase C synthesis (single claude-sonnet-4-5 call)...")
    print(f"{DIVIDER}\n")
    report = node5_synthesis.remote(mass_reactions, breakout_transcripts, brief)

    print(f"\n{'=' * 70}")
    print("  PANTHEON RESEARCH INTELLIGENCE REPORT")
    print(f"{'=' * 70}\n")
    print(report)
    print(f"\n{'=' * 70}")
    print("  END OF REPORT — PANTHEON EXECUTION COMPLETE")
    print(f"{'=' * 70}\n")

    # ── Save report locally ─────────────────────────────────────────────────
    md_path, base_name = _save_report(report, target, brief, client)

    # ── Node 6: Presentation Architect ──────────────────────────────────────
    _save_presentation(md_path, base_name, target, brief, client)

    # ── Node 7: Client Whisperer ────────────────────────────────────────────
    _run_client_whisperer(md_path, base_name, target, brief, client)

    return {
        "status": "success",
        "report": report,
        "agents_simulated": len(dynamic_agents),
        "breakout_groups": len(breakout_transcripts)
    }

# ─────────────────────────────────────────────────────────────────────────────
# FASTAPI SERVER & MODAL ASGI WRAPPER
# ─────────────────────────────────────────────────────────────────────────────

from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

web_app = FastAPI(title="PANTHEON API")

web_app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class PipelineRequest(BaseModel):
    target: str = "Medanese Upper Middle Class, 25-45"
    brief: str
    client: str = ""
    limit: int = 10
    group_size: int = 5

@web_app.post("/run_pipeline")
def api_run_pipeline(req: PipelineRequest):
    result = run_pipeline_core(
        target=req.target,
        brief=req.brief,
        client=req.client,
        limit=req.limit,
        group_size=req.group_size
    )
    return result

@app.function(image=image, secrets=SECRETS, timeout=3600)
@modal.asgi_app()
def fastapi_app():
    return web_app


# ─────────────────────────────────────────────────────────────────────────────
# NODE 6  — Presentation Architect
# Runs locally after Node 5 synthesis. Builds a pptxgenjs slide deck from the
# PANTHEON .md report using Claude Sonnet for content extraction.
# ─────────────────────────────────────────────────────────────────────────────

_DECK_CONTENT_TOOL = {
    "name": "extract_deck_content",
    "description": (
        "Extract structured content from a PANTHEON research report to populate "
        "a professional MBB-style slide deck. Distill each section to clear, "
        "concise language suitable for executive presentation."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "title": {
                "type": "string",
                "description": "Compelling deck title, e.g. 'BNPL Fintech: Medanese Consumer Insights'",
            },
            "brief_synopsis": {
                "type": "string",
                "description": "One sentence describing the campaign brief (≤20 words)",
            },
            "executive_summary": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string", "description": "Single bold headline (≤15 words)"},
                    "findings": {
                        "type": "array",
                        "items": {"type": "string", "description": "One finding (≤25 words)"},
                        "minItems": 3,
                        "maxItems": 5,
                    },
                },
                "required": ["headline", "findings"],
            },
            "market_context": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "stats": {
                        "type": "array",
                        "items": {"type": "string", "description": "Short stat or data point (≤15 words)"},
                        "minItems": 2,
                        "maxItems": 3,
                    },
                    "paragraph": {"type": "string", "description": "2-3 sentence context paragraph"},
                },
                "required": ["headline", "stats", "paragraph"],
            },
            "audience_insights": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "archetypes": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "name": {"type": "string"},
                                "age_range": {"type": "string"},
                                "profile": {"type": "string", "description": "2-3 sentences"},
                            },
                            "required": ["name", "age_range", "profile"],
                        },
                        "minItems": 2,
                        "maxItems": 3,
                    },
                },
                "required": ["headline", "archetypes"],
            },
            "consumer_response": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "findings": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "metric": {"type": "string", "description": "Short label (≤6 words)"},
                                "value": {"type": "string", "description": "Key metric value e.g. '72%' or 'High'"},
                                "detail": {"type": "string", "description": "One sentence explanation"},
                            },
                            "required": ["metric", "value", "detail"],
                        },
                        "minItems": 2,
                        "maxItems": 4,
                    },
                },
                "required": ["headline", "findings"],
            },
            "key_insights": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "insights": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "title": {"type": "string", "description": "Insight title (≤8 words)"},
                                "body": {"type": "string", "description": "1-2 sentence explanation"},
                            },
                            "required": ["title", "body"],
                        },
                        "minItems": 2,
                        "maxItems": 4,
                    },
                },
                "required": ["headline", "insights"],
            },
            "recommendations": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "priority": {"type": "string", "enum": ["High", "Medium", "Low"]},
                                "title": {"type": "string", "description": "Recommendation title (≤8 words)"},
                                "action": {"type": "string", "description": "One sentence action item"},
                            },
                            "required": ["priority", "title", "action"],
                        },
                        "minItems": 2,
                        "maxItems": 4,
                    },
                },
                "required": ["headline", "items"],
            },
            "risks": {
                "type": "object",
                "properties": {
                    "headline": {"type": "string"},
                    "items": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "level": {"type": "string", "enum": ["High", "Medium", "Low"]},
                                "title": {"type": "string", "description": "Risk title (≤8 words)"},
                                "indicator": {"type": "string", "description": "One sentence kill switch signal"},
                            },
                            "required": ["level", "title", "indicator"],
                        },
                        "minItems": 2,
                        "maxItems": 4,
                    },
                },
                "required": ["headline", "items"],
            },
        },
        "required": [
            "title", "brief_synopsis",
            "executive_summary", "market_context", "audience_insights",
            "consumer_response", "key_insights", "recommendations", "risks",
        ],
    },
}

# Node.js pptxgenjs deck template (MBB dark-navy design system)
_PPTXGENJS_TEMPLATE = r"""
const pptxgen = require('D:/npm-global/node_modules/pptxgenjs');
const fs = require('fs');

const data = JSON.parse(fs.readFileSync(process.argv[2], 'utf8'));
const outFile = process.argv[3];

const NAVY="1E2761", DARK="0D1B2A", ICE="CADCFC", GOLD="E8B04B";
const WHITE="F7F9FC", MUTED="8896A8", BODY="334155";

let pres = new pptxgen();
pres.layout = 'LAYOUT_16x9';
pres.author = 'PANTHEON';
pres.title = data.title;

function hdr(slide, title) {
    slide.addShape(pres.shapes.RECTANGLE, {x:0,y:0,w:10,h:0.72,fill:{color:NAVY},line:{color:NAVY}});
    slide.addShape(pres.shapes.RECTANGLE, {x:0,y:0,w:0.2,h:5.625,fill:{color:GOLD},line:{color:GOLD}});
    slide.addText(title, {x:0.4,y:0.1,w:9.3,h:0.52,fontSize:22,color:WHITE,bold:true,fontFace:"Calibri",valign:"middle",margin:0});
}

// §0 COVER
let s0 = pres.addSlide();
s0.background = {color:DARK};
s0.addShape(pres.shapes.RECTANGLE, {x:0,y:0,w:0.2,h:5.625,fill:{color:GOLD},line:{color:GOLD}});
s0.addText("PANTHEON\u2122", {x:7.0,y:0.22,w:2.8,h:0.38,fontSize:11,color:GOLD,bold:true,align:"right",fontFace:"Calibri"});
s0.addText(data.title, {x:0.45,y:1.0,w:9.3,h:1.55,fontSize:36,color:WHITE,bold:true,fontFace:"Calibri",valign:"middle"});
s0.addShape(pres.shapes.RECTANGLE, {x:0.45,y:2.72,w:2.8,h:0.045,fill:{color:GOLD},line:{color:GOLD}});
s0.addText(data.brief_synopsis, {x:0.45,y:2.87,w:9.2,h:0.78,fontSize:15,color:ICE,italic:true,fontFace:"Calibri",valign:"top"});
s0.addText("Target Demographic: " + data.target, {x:0.45,y:3.78,w:8.5,h:0.35,fontSize:11,color:MUTED});
s0.addText(data.date, {x:7.2,y:5.12,w:2.6,h:0.3,fontSize:10,color:MUTED,align:"right"});
s0.addText("Synthetic Focus Group Research", {x:0.45,y:5.12,w:5,h:0.3,fontSize:10,color:MUTED});

// §1 EXECUTIVE SUMMARY
let s1 = pres.addSlide();
s1.background = {color:WHITE};
hdr(s1, "Executive Summary");
s1.addText(data.executive_summary.headline, {x:0.45,y:0.82,w:9.3,h:0.42,fontSize:15,color:NAVY,bold:true,italic:true,fontFace:"Calibri"});
(data.executive_summary.findings || []).slice(0,5).forEach((f, i) => {
    let y = 1.38 + i * 0.75;
    s1.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y+0.04,w:0.06,h:0.4,fill:{color:GOLD},line:{color:GOLD}});
    s1.addText(f, {x:0.62,y:y,w:9.0,h:0.56,fontSize:13,color:BODY,fontFace:"Calibri",valign:"middle"});
});

// §2 AGENDA
let s2 = pres.addSlide();
s2.background = {color:WHITE};
hdr(s2, "Agenda");
["Market & Demographic Context","Audience Psychology & Deep-Dive","Consumer Response Analysis","Key Insights & Patterns","Strategic Recommendations","Risk Assessment & Kill Switch Signals"].forEach((sec, i) => {
    let y = 0.9 + i * 0.72;
    let nc = i % 2 === 0 ? GOLD : NAVY;
    s2.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y+0.04,w:0.34,h:0.34,fill:{color:nc},line:{color:nc}});
    s2.addText(String(i+1), {x:0.45,y:y+0.04,w:0.34,h:0.34,fontSize:12,color:WHITE,bold:true,align:"center",valign:"middle",margin:0});
    s2.addText(sec, {x:0.92,y:y,w:8.8,h:0.44,fontSize:14,color:BODY,fontFace:"Calibri",valign:"middle"});
});

// §3 MARKET CONTEXT
let s3 = pres.addSlide();
s3.background = {color:WHITE};
hdr(s3, "Market & Demographic Context");
s3.addText(data.market_context.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
(data.market_context.stats || []).slice(0,3).forEach((stat, i) => {
    let x = 0.45 + i * 3.12;
    s3.addShape(pres.shapes.RECTANGLE, {x:x,y:1.35,w:2.92,h:1.38,fill:{color:NAVY},line:{color:NAVY}});
    s3.addText(stat, {x:x+0.1,y:1.38,w:2.72,h:1.32,fontSize:13,color:ICE,fontFace:"Calibri",valign:"middle",align:"center"});
});
s3.addText(data.market_context.paragraph, {x:0.45,y:2.9,w:9.3,h:2.45,fontSize:12,color:BODY,fontFace:"Calibri",valign:"top"});

// §4 AUDIENCE
let s4 = pres.addSlide();
s4.background = {color:WHITE};
hdr(s4, "Audience Psychology & Deep-Dive");
s4.addText(data.audience_insights.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
let archs = (data.audience_insights.archetypes || []).slice(0,3);
let cw = archs.length > 0 ? (9.1 / archs.length) - 0.1 : 9.1;
archs.forEach((arch, i) => {
    let x = 0.45 + i * (cw + 0.1);
    s4.addShape(pres.shapes.RECTANGLE, {x:x,y:1.35,w:cw,h:3.98,fill:{color:"F0F4F8"},line:{color:"E2E8F0"}});
    s4.addShape(pres.shapes.RECTANGLE, {x:x,y:1.35,w:cw,h:0.32,fill:{color:NAVY},line:{color:NAVY}});
    s4.addText(arch.name, {x:x+0.08,y:1.35,w:cw-0.16,h:0.32,fontSize:11,color:WHITE,bold:true,valign:"middle",margin:0});
    s4.addText("Age: " + (arch.age_range || ""), {x:x+0.08,y:1.76,w:cw-0.16,h:0.3,fontSize:10,color:GOLD,bold:true});
    s4.addText(arch.profile || "", {x:x+0.08,y:2.12,w:cw-0.16,h:3.1,fontSize:11,color:BODY,fontFace:"Calibri",valign:"top"});
});

// §5 CONSUMER RESPONSE
let s5 = pres.addSlide();
s5.background = {color:WHITE};
hdr(s5, "Consumer Response Analysis");
s5.addText(data.consumer_response.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
(data.consumer_response.findings || []).slice(0,4).forEach((f, i) => {
    let y = 1.35 + i * 0.97;
    s5.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:9.3,h:0.87,fill:{color:"F8FAFC"},line:{color:"E2E8F0"}});
    s5.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:0.06,h:0.87,fill:{color:GOLD},line:{color:GOLD}});
    s5.addText(f.metric || "", {x:0.62,y:y+0.06,w:1.8,h:0.32,fontSize:10,color:MUTED,bold:true});
    s5.addText(f.value || "", {x:2.55,y:y+0.04,w:1.2,h:0.38,fontSize:18,color:NAVY,bold:true,align:"center"});
    s5.addText(f.detail || "", {x:0.62,y:y+0.44,w:8.9,h:0.36,fontSize:11,color:BODY});
});

// §6 KEY INSIGHTS
let s6 = pres.addSlide();
s6.background = {color:WHITE};
hdr(s6, "Key Insights & Patterns");
s6.addText(data.key_insights.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
(data.key_insights.insights || []).slice(0,4).forEach((ins, i) => {
    let y = 1.35 + i * 0.97;
    s6.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:9.3,h:0.87,fill:{color:"F8FAFC"},line:{color:"E2E8F0"}});
    s6.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:0.06,h:0.87,fill:{color:GOLD},line:{color:GOLD}});
    s6.addText(ins.title || "", {x:0.62,y:y+0.06,w:9.0,h:0.3,fontSize:12,color:NAVY,bold:true});
    s6.addText(ins.body || "", {x:0.62,y:y+0.41,w:9.0,h:0.4,fontSize:11,color:BODY});
});

// §7 RECOMMENDATIONS
let s7 = pres.addSlide();
s7.background = {color:WHITE};
hdr(s7, "Strategic Recommendations");
s7.addText(data.recommendations.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
(data.recommendations.items || []).slice(0,4).forEach((rec, i) => {
    let y = 1.35 + i * 0.97;
    let pc = rec.priority === "High" ? "DC2626" : (rec.priority === "Medium" ? GOLD : "22C55E");
    s7.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:9.3,h:0.87,fill:{color:"F8FAFC"},line:{color:"E2E8F0"}});
    s7.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:0.06,h:0.87,fill:{color:pc},line:{color:pc}});
    s7.addShape(pres.shapes.RECTANGLE, {x:8.52,y:y+0.19,w:0.88,h:0.27,fill:{color:pc},line:{color:pc}});
    s7.addText(rec.priority || "", {x:8.52,y:y+0.19,w:0.88,h:0.27,fontSize:9,color:WHITE,bold:true,align:"center",valign:"middle",margin:0});
    s7.addText(rec.title || "", {x:0.62,y:y+0.06,w:7.7,h:0.3,fontSize:12,color:NAVY,bold:true});
    s7.addText(rec.action || "", {x:0.62,y:y+0.42,w:9.0,h:0.38,fontSize:11,color:BODY});
});

// §8 RISKS
let s8 = pres.addSlide();
s8.background = {color:WHITE};
hdr(s8, "Risk Assessment & Kill Switch Signals");
s8.addText(data.risks.headline, {x:0.45,y:0.82,w:9.3,h:0.4,fontSize:15,color:NAVY,bold:true,italic:true});
(data.risks.items || []).slice(0,4).forEach((r, i) => {
    let y = 1.35 + i * 0.97;
    let rc = r.level === "High" ? "DC2626" : (r.level === "Medium" ? GOLD : "22C55E");
    s8.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:9.3,h:0.87,fill:{color:"F8FAFC"},line:{color:"E2E8F0"}});
    s8.addShape(pres.shapes.RECTANGLE, {x:0.45,y:y,w:0.06,h:0.87,fill:{color:rc},line:{color:rc}});
    s8.addShape(pres.shapes.RECTANGLE, {x:8.52,y:y+0.19,w:0.88,h:0.27,fill:{color:rc},line:{color:rc}});
    s8.addText(r.level || "", {x:8.52,y:y+0.19,w:0.88,h:0.27,fontSize:9,color:WHITE,bold:true,align:"center",valign:"middle",margin:0});
    s8.addText(r.title || "", {x:0.62,y:y+0.06,w:7.7,h:0.3,fontSize:12,color:NAVY,bold:true});
    s8.addText(r.indicator || "", {x:0.62,y:y+0.42,w:9.0,h:0.38,fontSize:11,color:BODY});
});

// §9 CLOSING
let s9 = pres.addSlide();
s9.background = {color:DARK};
s9.addShape(pres.shapes.RECTANGLE, {x:0,y:0,w:0.2,h:5.625,fill:{color:GOLD},line:{color:GOLD}});
s9.addText("PANTHEON\u2122", {x:0.45,y:1.1,w:9.1,h:1.05,fontSize:52,color:WHITE,bold:true,fontFace:"Calibri",align:"center"});
s9.addShape(pres.shapes.RECTANGLE, {x:3.1,y:2.25,w:3.8,h:0.045,fill:{color:GOLD},line:{color:GOLD}});
s9.addText("Synthetic Focus Group Research Engine", {x:0.45,y:2.38,w:9.1,h:0.44,fontSize:14,color:ICE,italic:true,align:"center"});
s9.addText("Powered by Anthropic Claude \u00b7 Modal Serverless \u00b7 Supabase", {x:0.45,y:2.95,w:9.1,h:0.35,fontSize:11,color:MUTED,align:"center"});
s9.addText("This report was generated using synthetic focus group simulation. All respondents are AI-synthesized personas. Results are for strategic guidance only.", {x:1.2,y:4.55,w:7.6,h:0.72,fontSize:9,color:MUTED,align:"center",italic:true});
s9.addText(data.date, {x:7.2,y:5.12,w:2.6,h:0.3,fontSize:10,color:MUTED,align:"right"});

pres.writeFile({fileName: outFile})
    .then(() => { console.log("PPTX_DONE"); process.exit(0); })
    .catch(e => { console.error("PPTX_ERR: " + e.message); process.exit(1); });
"""


def _extract_deck_content(
    anthropic_client,
    md_content: str,
    target: str,
    brief: str,
    client: str,
) -> dict | None:
    """
    Call Claude Sonnet (tool-forced) to extract structured deck content
    from the PANTHEON .md report. Returns the extracted JSON dict or None.
    """
    import anthropic as _anthropic

    system = (
        "You are an expert MBB (McKinsey/Bain/BCG) presentation strategist. "
        "Extract concise, punchy content from the PANTHEON research report for a "
        "professional executive slide deck. Be specific and data-driven. "
        "Every field must be populated with meaningful content from the report."
    )
    user_msg = (
        f"Extract slide deck content from this PANTHEON Research Intelligence Report.\n\n"
        f"Client/Product context: {client or 'Not specified'}\n"
        f"Target demographic: {target}\n"
        f"Campaign brief: {brief}\n\n"
        f"--- REPORT ---\n{md_content[:12000]}"
    )

    for attempt in range(1, 3):
        try:
            resp = anthropic_client.messages.create(
                model="claude-sonnet-4-5",
                max_tokens=4096,
                system=system,
                tools=[_DECK_CONTENT_TOOL],
                tool_choice={"type": "tool", "name": "extract_deck_content"},
                messages=[{"role": "user", "content": user_msg}],
            )
            block = next((b for b in resp.content if b.type == "tool_use"), None)
            if block is None:
                raise ValueError("No tool_use block returned")
            return block.input
        except _anthropic.RateLimitError:
            print("  Node 6: rate limit hit — sleeping 65s...")
            time.sleep(65)
        except Exception as e:
            print(f"  Node 6: extraction attempt {attempt}/2 failed: {e}")
            if attempt < 2:
                time.sleep(5)
    return None


def _save_presentation(md_path: Path, base_name: str, target: str, brief: str, client: str = "") -> None:
    """
    Node 6 — Presentation Architect.
    Reads the PANTHEON .md report, extracts structured content via Claude Sonnet,
    generates a pptxgenjs slide deck (.pptx), and emits a PANTHEON_PPTX_FILE:: sentinel.
    Always runs on the local machine (no Modal container).
    """
    print(f"\n{DIVIDER}")
    print("  Node 6 — Presentation Architect (pptxgenjs · 10-slide deck)")
    print(DIVIDER)

    try:
        md_content = md_path.read_text(encoding="utf-8")
    except Exception as e:
        print(f"  Node 6: Cannot read .md report: {e} — skipping presentation.")
        return

    # ── Step 1: Extract structured content via Claude Sonnet ─────────────────
    from dotenv import load_dotenv
    import anthropic as _anthropic

    load_dotenv(Path(__file__).parent / "pantheon.env", override=True)
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    ac = _anthropic.Anthropic(api_key=api_key)

    print("  Extracting structured deck content via Claude Sonnet...")
    deck_data = _extract_deck_content(ac, md_content, target, brief, client)
    if not deck_data:
        print("  Node 6: Content extraction failed — skipping presentation.")
        return

    # Inject system-controlled fields
    deck_data["target"] = target
    deck_data["date"] = datetime.now().strftime("%B %Y")

    # ── Step 2: Write JSON data file & JS script ──────────────────────────────
    out_dir = md_path.parent
    json_file = f"{base_name}_deck_data.json"
    js_file   = f"{base_name}_gen.js"
    pptx_file = f"{base_name}.pptx"

    json_path = out_dir / json_file
    js_path   = out_dir / js_file
    pptx_path = out_dir / pptx_file

    json_path.write_text(json.dumps(deck_data, ensure_ascii=False, indent=2), encoding="utf-8")
    js_path.write_text(_PPTXGENJS_TEMPLATE, encoding="utf-8")

    # ── Step 3: Run Node.js to generate the PPTX ─────────────────────────────
    try:
        import subprocess
        result = subprocess.run(
            ["node", js_file, json_file, pptx_file],
            capture_output=True,
            text=True,
            timeout=90,
            cwd=str(out_dir),
        )
        if result.returncode != 0 or not pptx_path.exists():
            err = (result.stderr or result.stdout or "Unknown error").strip()
            print(f"  Node 6: pptxgenjs failed:\n    {err}")
            return
        size_kb = pptx_path.stat().st_size // 1024
        print(f"  Deck saved ({size_kb} KB) → {pptx_path}")
        print(f"PANTHEON_PPTX_FILE::{base_name}")
    except subprocess.TimeoutExpired:
        print("  Node 6: Node.js timed out after 90s — skipping presentation.")
    except Exception as e:
        print(f"  Node 6: Execution error: {e}")
    finally:
        # Clean up temp files
        for tmp in [json_path, js_path]:
            try:
                tmp.unlink(missing_ok=True)
            except Exception:
                pass


def _save_report(report: str, target: str, brief: str, client: str = "") -> tuple[Path, str]:
    """
    Save the Phase C report to disk as both a Word document (.docx) and
    a plain Markdown file (.md). Always runs on the local machine.
    Returns (md_path, base_name).
    """
    # Derive a filesystem-safe slug from the target demographic
    slug = re.sub(r"[^\w]+", "_", target).strip("_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")

    # ── Client-organised output folder ──────────────────────────────────────
    client_folder = re.sub(r"[^\w]+", "_", client).strip("_") if client.strip() else "Unnamed"
    out_dir = Path(__file__).parent / "reports" / client_folder
    out_dir.mkdir(parents=True, exist_ok=True)
    print(f"PANTHEON_CLIENT_FOLDER::{client_folder}")  # sentinel for dashboard

    if client:
        client_slug = re.sub(r"[^\w]+", "_", client).strip("_")
        version = 1
        while True:
            base_name = f"PANTHEON_REPORT_{client_slug}_v{version}"
            if not ((out_dir / f"{base_name}.docx").exists() or (out_dir / f"{base_name}.md").exists()):
                break
            version += 1
    else:
        base_name = f"PANTHEON_Report_{slug}"

    docx_path = out_dir / f"{base_name}.docx"
    md_path   = out_dir / f"{base_name}.md"

    print(f"PANTHEON_OUTPUT_FILE::{base_name}")

    # ── Markdown fallback (always succeeds) ───────────────────────────────
    md_content = (
        f"# PANTHEON Research Intelligence Report\n"
        f"**Target demographic:** {target}  \n"
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  \n"
        f"**Brief:** {brief}\n\n"
        f"---\n\n"
        + report
    )
    md_path.write_text(md_content, encoding="utf-8")
    print(f"  Markdown saved  → {md_path}")

    # ── Word document ─────────────────────────────────────────────────────
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # --- Title ---
        title_para = doc.add_heading("PANTHEON Research Intelligence Report", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- Metadata block ---
        meta = doc.add_paragraph()
        meta.add_run("Target demographic: ").bold = True
        meta.add_run(target)
        meta.add_run("\nGenerated: ").bold = True
        meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        meta.add_run("\nCampaign brief: ").bold = True
        meta.add_run(brief)
        doc.add_paragraph()   # spacer

        # --- Parse and render each line of the report ---
        # Strategy: detect markdown-style headers (##, **bold**) and
        # render them as Word headings / bold runs; everything else is body.
        for raw_line in report.splitlines():
            line = raw_line.rstrip()

            # Blank lines → paragraph spacer
            if not line.strip():
                doc.add_paragraph()
                continue

            # h2 → Word Heading 2
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                continue

            # h3 → Word Heading 3
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                continue

            # h1/h4+ → Word Heading 1
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
                continue

            # Horizontal rule (---, ===, ───)
            if re.match(r"^[-=─]{3,}\s*$", line):
                doc.add_paragraph("─" * 50)
                continue

            # Normal paragraph — handle inline **bold** markers
            para = doc.add_paragraph()
            # Split on **...** pairs
            segments = re.split(r"(\*\*[^*]+\*\*)", line)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    run = para.add_run(seg[2:-2])
                    run.bold = True
                else:
                    para.add_run(seg)

        doc.save(str(docx_path))
        print(f"  Word doc saved  → {docx_path}")

    except ImportError:
        print("  python-docx not installed — skipping .docx output (pip install python-docx)")
    except Exception as exc:
        print(f"  Word doc failed : {exc}")
        print(f"  (Markdown fallback at {md_path} is intact)")

    return md_path, base_name


# ─────────────────────────────────────────────────────────────────────────────
# NODE 7  — Client Whisperer
# Runs locally after Node 6. Translates the PANTHEON .md report into a 
# strategically framed, empathy-driven meeting preparation document in .docx.
# Built for client-facing teams to steer complex strategic conversations.
# ─────────────────────────────────────────────────────────────────────────────

def _run_client_whisperer(md_path: Path, base_name: str, target: str, brief: str, client: str = "") -> None:
    """
    Node 7 — Client Whisperer.
    Distills PANTHEON research and context into a client-facing meeting prep doc.
    Outputs a formatted .docx file.
    """
    print(f"\n{DIVIDER}")
    print("  Node 7 — Client Whisperer (Strategic Meeting Prep)")
    print(DIVIDER)

    try:
        md_content = md_path.read_text(encoding="utf-8")
    except Exception as e:
        print(f"  Node 7: Cannot read .md report: {e} — skipping Whisperer.")
        return

    from dotenv import load_dotenv
    import anthropic as _anthropic

    load_dotenv(Path(__file__).parent / "pantheon.env", override=True)
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    ac = _anthropic.Anthropic(api_key=api_key)

    system_prompt = """You are the Client Whisperer — Storytellers' most important front-facing
intelligence. You are not a salesperson. You are a trusted advisor who
happens to represent a firm that can solve what you're about to uncover.
Your role is to take the full output of the pipeline — PANTHEON's research,
the Presentation Architect's deck — and distill it into a human conversation.
You translate data into empathy. You translate insights into questions.
You translate recommendations into a path forward that makes the client
feel understood before they ever feel sold to.
You think like a strategist. You speak like a person who genuinely cares.
You close like someone who already knows the answer — because you do.

═══════════════════════════════════════════════════════════════
PART I — BEFORE ANYTHING ELSE: SANITY CHECK
═══════════════════════════════════════════════════════════════
MANDATORY FIRST STEP — DO NOT SKIP:
Before preparing any client-facing material, you must verify that the
services being recommended are within Storytellers' actual capability.
Run the following internal check against every recommendation you intend
to make:

STORYTELLERS SERVICE SCOPE (as understood from brief):
Storytellers Creative Solutions / Storytellers Asia is a strategic
marketing advisory and creative agency. Core services include:

Marketing strategy & campaign development
Brand strategy, positioning, and identity
Creative direction and content production
Performance marketing and digital strategy
Go-to-market strategy and launch planning
CRM strategy and first-party data planning
Revenue growth strategy and monetization consulting

FOR EACH RECOMMENDATION, ASK:
[?] Is this something Storytellers can execute or advise on?
[?] Is this within the strategic advisory scope?
[?] Does this require a capability Storytellers does not have?

FLAG SYSTEM:
[✓ IN SCOPE] — Storytellers can deliver this
[~ ADJACENT] — Storytellers can advise; execution partner may be needed
[✗ OUT OF SCOPE] — Do not pitch this; refer out or exclude

RULE: Never present a recommendation or CTA for a service that has not
cleared [✓ IN SCOPE] or [~ ADJACENT] status.
If a core finding demands a solution that is [✗ OUT OF SCOPE],
acknowledge the problem honestly and recommend the appropriate
external resource — then identify the Storytellers-adjacent angle.

═══════════════════════════════════════════════════════════════
PART II — INPUT PROCESSING
═══════════════════════════════════════════════════════════════

PARSING SEQUENCE — 4 PASSES:
PASS 1: BUSINESS READING
Extract: What is this company? What category? What stage?
What do their materials say about how they see themselves?
What do they say their problem is?

PASS 2: BRAND DECODING
From all available materials, reconstruct:
BRAND_VOICE: How do they speak? (3 tone descriptors)
BRAND_VALUES: What do they stand for?
BRAND_GAP: Where does the promise diverge from reality?
VISUAL_LANG: Aesthetic register?
COMM_PATTERN: Past communication style?

PASS 3: PAIN EXTRACTION
From findings, identify:
CORE_STRUGGLE: Central business/marketing problem
SURFACE_PAIN: What they say the problem is
REAL_PAIN: What data reveals the problem actually is
UNSPOKEN_FEAR: What they fear saying out loud
CONSEQUENCE: What happens if unfixed
PRIDE_POINT: What they are proud of (never attack)
SENSITIVITY_ZONE: Trigger topics

PASS 4: SOLUTION MAPPING
Map each pain to a Storytellers service (post-sanity check):
PAIN → ROOT_CAUSE → STORYTELLERS_LEVER → EXPECTED_OUTCOME
Identify 1 urgent anchor, 2 supporting expansion points (Max 3).

═══════════════════════════════════════════════════════════════
PART III — MEETING NOTES OUTPUT FORMAT (Markdown)
═══════════════════════════════════════════════════════════════
Output your meeting prep document as raw Markdown text. Do NOT use markdown code blocks like ```markdown ..., just the raw content.
Use standard `# Heading 1`, `## Heading 2`, and `**bolding**`.

DOCUMENT STRUCTURE:
────────────────────────────────────────
SECTION 1: CLIENT SNAPSHOT (internal only)
────────────────────────────────────────
A rapid-read summary for whoever walks into this meeting.
Company: [name + category + stage]
Core Business: [1-line description]
Brand Read: [voice + values + gap]
Market Position: [leader/challenger/etc.]
Key Tension: [single most important thing to understand]
What They Think: [stated problem]
What's Real: [actual problem]
Pride Point: [never attack this]
Sensitivity Zone: [proceed carefully]

────────────────────────────────────────
SECTION 2: THE CONVERSATION ARCHITECTURE
────────────────────────────────────────
STAGE 1: ESTABLISH (5–10 min)
Opening statement (NOT a question) demonstrating understanding.

STAGE 2: PROBE (15–20 min)
Questions opening the wound. Sequence: aspiration -> friction -> consequence -> emotion -> ownership.
Format:
Q[N] [QUESTION TEXT]
PURPOSE: [what this surfaces]
EXPECTED RESPONSE TERRITORY: [typical answer]
FOLLOW-UP IF THEY OPEN UP: [go deeper]
BACK-OUT IF THEY CLOSE: [redirect]

STAGE 3: REFLECT (5 min)
Mirroring script template.

STAGE 4: REFRAME (5–10 min)
Reframe template: "Most businesses think the problem is [SURFACE]. What we find is [REAL]."

STAGE 5: FRAMEWORK (5–10 min)
Maximum 4-step solution map in plain language.
Each step: WHAT WE DO → WHY IT MATTERS → WHAT CHANGES

STAGE 6: CTA (5 min)
One clear next step (LOW, MED, or HIGH friction).

────────────────────────────────────────
SECTION 3: SIGNAL READING GUIDE
────────────────────────────────────────
OPEN SIGNALS (proceed deeper)
CLOSE SIGNALS (back off, redirect)
BACK-OUT SCRIPTS (use when close signal detected)

────────────────────────────────────────
SECTION 4: PLAIN LANGUAGE TRANSLATION GUIDE
────────────────────────────────────────
Standard format:
PROFESSIONAL: [original]
PLAIN: [rewritten for grandmother]
ANALOGY: [relatable analogy]

────────────────────────────────────────
SECTION 5: STORYTELLERS CTA & SERVICE FIT
────────────────────────────────────────
Only IN SCOPE / ADJACENT services:
SERVICE: [Service name and flag]
PROBLEM IT SOLVES: [pain]
HOW TO INTRODUCE IT: [exact language]
PROOF POINT: [analogy/result]
WHAT WE'RE NOT: [boundary]
SERVICES EXCLUDED: [list any OUT OF SCOPE with reason]

────────────────────────────────────────
SECTION 6: MEETING LOGISTICS
────────────────────────────────────────
Duration, Attendees, Materials, Pre-meet action, Post-meet action.

═══════════════════════════════════════════════════════════════
PART IV & V — LANGUAGE & BEHAVIOR
═══════════════════════════════════════════════════════════════
- Internal sections are direct/analytical. Client-facing sections are warm/grounded.
- No jargon: NEVER say leverage, synergy, holistic.
- Always include the sanity check status for services.
"""

    user_msg = (
        f"Generate the Meeting Prep Document based on the following context.\n\n"
        f"Client/Product context: {client or '(Unknown/Generic)'}\n"
        f"Target demographic: {target}\n"
        f"Campaign brief: {brief}\n\n"
        f"--- PANTHEON RESEARCH INTELLIGENCE REPORT ---\n{md_content[:40000]}\n"
    )

    print("  Distilling insights via Claude Sonnet...")
    whisperer_md = ""
    for attempt in range(1, 3):
        try:
            resp = ac.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=8192,
                system=system_prompt,
                messages=[{"role": "user", "content": user_msg}],
            )
            whisperer_md = resp.content[0].text
            break
        except _anthropic.RateLimitError:
            print("  Node 7: rate limit hit — sleeping 65s...")
            time.sleep(65)
        except Exception as e:
            print(f"  Node 7: Generation attempt {attempt}/2 failed: {e}")
            time.sleep(5)
            
    if not whisperer_md:
        print("  Node 7: Failed to generate meeting prep document.")
        return

    # Clean off potential markdown codeblock formatting if Claude added it
    whisperer_md = re.sub(r"^```(?:markdown)?\n", "", whisperer_md, flags=re.IGNORECASE).strip()
    whisperer_md = re.sub(r"\n```$", "", whisperer_md).strip()

    # ── Write to docx ────────────────────────────────────────────────────────
    client_safe = re.sub(r"[^\w]+", "", client) if client else "UntitledClient"
    date_str = datetime.now().strftime("%Y%m%d")
    out_dir = md_path.parent
    
    # Prepend PANTHEON base name to keep directory tidy
    docx_file_name = f"{base_name}_ClientWhisperer_{client_safe}_{date_str}.docx"
    docx_path = out_dir / docx_file_name

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        
        # Parse and render Markdown to docx (simple version)
        for raw_line in whisperer_md.splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                doc.add_paragraph()
                continue
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                continue
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
                continue
            if re.match(r"^[-=─]{3,}\s*$", line):
                doc.add_paragraph("─" * 50)
                continue

            para = doc.add_paragraph()
            segments = re.split(r"(\*\*[^*]+\*\*)", line)
            for seg in segments:
                if seg.startswith("**") and seg.endswith("**"):
                    run = para.add_run(seg[2:-2])
                    run.bold = True
                else:
                    para.add_run(seg)

        doc.save(str(docx_path))
        print(f"  Meeting Prep saved ({docx_path.stat().st_size // 1024} KB) → {docx_path}")
        print(f"PANTHEON_WHISPERER_FILE::{docx_file_name}")
    except ImportError:
        print("  Node 7: python-docx not installed. Run 'pip install python-docx'.")
    except Exception as e:
        print(f"  Node 7: Failed to save .docx: {e}")
        # Fallback to .md
        fallback_file = f"{base_name}_ClientWhisperer_{client_safe}_{date_str}.md"
        fallback_path = out_dir / fallback_file
        fallback_path.write_text(whisperer_md, encoding="utf-8")
        print(f"  Saved fallback Markdown → {fallback_path}")
        print(f"PANTHEON_WHISPERER_FILE::{fallback_file}")
